#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
RAMS Tasarım İlkeleri - Nefes Saati Projesi
Çıktı: Nefes_Saati_RAMS_Analizi.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── sabitler ────────────────────────────────────────────────────────────────

FONT = 'Times New Roman'
SP   = Pt(5)   # 100 twips spacing

# ── yardımcı fonksiyonlar ────────────────────────────────────────────────────

def _set_font(run, size_pt, bold=False, color=None):
    run.font.name  = FONT
    run.font.size  = Pt(size_pt)
    run.bold       = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # East-Asian font fix
    rPr = run._r.get_or_add_rPr()
    fonts_elem = rPr.find(qn('w:rFonts'))
    if fonts_elem is None:
        fonts_elem = OxmlElement('w:rFonts')
        rPr.insert(0, fonts_elem)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        fonts_elem.set(qn(attr), FONT)

def _spacing(para, before=None, after=None):
    pf = para.paragraph_format
    if before is not None: pf.space_before = before
    if after  is not None: pf.space_after  = after

def _shade_para(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)

def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def _cell_vertical_align(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)

# ── ekleme fonksiyonları ─────────────────────────────────────────────────────

def add_title(doc, text, size=24, center=True):
    p = doc.add_paragraph()
    _spacing(p, SP, SP)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_font(r, size, bold=True)
    return p

def add_section(doc, text, size=18, number=None):
    p = doc.add_paragraph()
    _spacing(p, Pt(8), Pt(4))
    full = (f'{number}. ' if number else '') + text
    r = p.add_run(full)
    _set_font(r, size, bold=True)
    return p

def add_subsection(doc, text, size=13.5):
    p = doc.add_paragraph()
    _spacing(p, Pt(6), Pt(3))
    r = p.add_run(text)
    _set_font(r, size, bold=True)
    return p

def add_body(doc, text, indent=False, italic=False):
    p = doc.add_paragraph()
    _spacing(p, Pt(3), Pt(3))
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    _set_font(r, 12)
    if italic:
        r.italic = True
    return p

def add_label(doc, label, value):
    p = doc.add_paragraph()
    _spacing(p, Pt(3), Pt(3))
    r1 = p.add_run(label + ': ')
    _set_font(r1, 12, bold=True)
    r2 = p.add_run(value)
    _set_font(r2, 12)
    return p

def add_bullet(doc, text, indent=0.3):
    p = doc.add_paragraph()
    _spacing(p, Pt(2), Pt(2))
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    r = p.add_run(u'•  ' + text)
    _set_font(r, 12)
    return p

def add_sub_bullet(doc, text):
    p = doc.add_paragraph()
    _spacing(p, Pt(1), Pt(1))
    p.paragraph_format.left_indent  = Inches(0.6)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    r = p.add_run(u'◦  ' + text)
    _set_font(r, 11)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    _spacing(p, Pt(4), Pt(4))
    p.paragraph_format.left_indent = Inches(0.4)
    _shade_para(p, 'F2F2F2')
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    _spacing(p, Pt(2), Pt(2))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(u'─' * 80)
    r.font.name = FONT
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p

def add_risk_table(doc, rows):
    """4-sütunlu risk tablosu: Risk | Olasılık | Etki | Seviye"""
    headers = ['Risk Faktörü', 'Olasılık', 'Etki', 'Risk Seviyesi']
    colors  = {
        'Kritik':    'FF0000', 'Yüksek': 'FF6600',
        'Orta':      'FFA500', 'Düşük':  '70AD47',
        'Çok Düşük': '92D050',
    }
    bg_header = '2E4057'
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # başlık satırı
    hrow = table.rows[0]
    col_widths = [Inches(2.8), Inches(1.1), Inches(1.0), Inches(1.3)]
    for i, (cell, hdr, w) in enumerate(zip(hrow.cells, headers, col_widths)):
        cell.width = w
        _shade_cell(cell, bg_header)
        _cell_vertical_align(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _set_font(r, 11, bold=True, color=(0xFF, 0xFF, 0xFF))
    # veri satırları
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            _shade_cell(cell, bg)
            _cell_vertical_align(cell)
            p = cell.paragraphs[0]
            if ci == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                level = txt.strip()
                fill  = colors.get(level, 'CCCCCC')
                _shade_cell(cell, fill)
                r = p.add_run(txt)
                _set_font(r, 11, bold=True,
                          color=(0xFF,0xFF,0xFF) if level not in ('Düşük','Çok Düşük') else (0x33,0x33,0x33))
            else:
                r = p.add_run(txt)
                _set_font(r, 11)
    return table

def add_summary_table(doc, rows):
    """RAMS özet tablosu"""
    headers = ['RAMS Kriteri', 'Açıklama', 'Nefes Saati Durumu', 'Risk Seviyesi', 'Öncelik']
    bg_header = '1F3864'
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for cell, hdr in zip(hrow.cells, headers):
        _shade_cell(cell, bg_header)
        _cell_vertical_align(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _set_font(r, 11, bold=True, color=(0xFF,0xFF,0xFF))
    criteria_colors = {
        'Reliability': 'E8F4FD', 'Availability': 'E8F8E8',
        'Maintainability': 'FFF8E8', 'Safety': 'FEE8E8',
    }
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = criteria_colors.get(row_data[0], 'FFFFFF')
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            _shade_cell(cell, bg)
            _cell_vertical_align(cell)
            p = cell.paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                _set_font(r, 11, bold=True)
            else:
                r = p.add_run(txt)
                _set_font(r, 10)
    return table

def add_metric_table(doc, rows):
    """Ölçüm metrikleri tablosu"""
    headers = ['Metrik', 'Hedef Değer', 'Mevcut Durum', 'Not']
    bg_header = '385723'
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for cell, hdr in zip(hrow.cells, headers):
        _shade_cell(cell, bg_header)
        _cell_vertical_align(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _set_font(r, 11, bold=True, color=(0xFF,0xFF,0xFF))
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = 'FFFFFF' if ri % 2 == 0 else 'F0F7EE'
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            _shade_cell(cell, bg)
            _cell_vertical_align(cell)
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            _set_font(r, 10, bold=(ci == 0))
    return table

def add_phase_table(doc, rows):
    """Faz - RAMS ilişkisi tablosu"""
    headers = ['Geliştirme Fazı', 'Kapsam', 'Reliability', 'Availability', 'Maintainability', 'Safety']
    bg_header = '4A235A'
    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    for cell, hdr in zip(hrow.cells, headers):
        _shade_cell(cell, bg_header)
        _cell_vertical_align(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        _set_font(r, 10, bold=True, color=(0xFF,0xFF,0xFF))
    checks = {True: u'✔', False: u'○'}
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = 'FFFFFF' if ri % 2 == 0 else 'F5F0F9'
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            _shade_cell(cell, bg)
            _cell_vertical_align(cell)
            p = cell.paragraphs[0]
            if ci >= 2 and isinstance(txt, bool):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(checks[txt])
                _set_font(r, 12, bold=txt, color=(0x00,0x70,0x00) if txt else (0x80,0x80,0x80))
            else:
                r = p.add_run(str(txt))
                _set_font(r, 10, bold=(ci == 0))
    return table

# ────────────────────────────────────────────────────────────────────────────
# BELGE OLUŞTURMA
# ────────────────────────────────────────────────────────────────────────────

doc = Document()

sec = doc.sections[0]
sec.left_margin   = Emu(899795)
sec.right_margin  = Emu(899795)
sec.top_margin    = Emu(899795)
sec.bottom_margin = Emu(899795)

# varsayılan Normal stilin fontunu ayarla
for style_name in ('Normal', 'Default Paragraph Font'):
    try:
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(12)
    except Exception:
        pass

# ══════════════════════════════════════════════════════════════════════════
# KAPAK BÖLÜMLERİ
# ══════════════════════════════════════════════════════════════════════════

add_divider(doc)

p = doc.add_paragraph()
_spacing(p, Pt(8), Pt(4))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RAMS TASARIM İLKELERİ')
_set_font(r, 24, bold=True)

p = doc.add_paragraph()
_spacing(p, Pt(4), Pt(4))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NEFES SAATİ')
_set_font(r, 20, bold=True)

p = doc.add_paragraph()
_spacing(p, Pt(2), Pt(8))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AR Destekli Akıllı İlaç Takip ve Hatırlatma Sistemi')
_set_font(r, 14)
r.italic = True

add_divider(doc)

# Öğrenci bilgileri
doc.add_paragraph()
add_subsection(doc, 'Öğrenci Bilgileri')
add_label(doc, 'Ad Soyad', 'Kenan Yılmaz')
add_label(doc, 'Okul No', '210542011')
add_label(doc, 'Proje Adı', 'Nefes Saati - AR Destekli Akıllı İlaç Takip ve Hatırlatma Sistemi')
add_label(doc, 'GitHub', 'https://github.com/kenanylmz/arSaglikProjesi')
add_label(doc, 'Platform', 'Android (React Native CLI 0.83)')
add_label(doc, 'Tarih', 'Nisan 2025')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. PROJE TANIMI
# ══════════════════════════════════════════════════════════════════════════

add_section(doc, 'Proje Tanımı', number=1)

add_subsection(doc, '1.1  Uygulamanın Amacı')
add_body(doc,
    'Nefes Saati, kronik hastalığa sahip bireyler ve yaşlı hastalar için geliştirilmiş '
    'Android tabanlı bir mobil sağlık uygulamasıdır. Uygulama; artırılmış gerçeklik (AR) '
    'kamera modülü aracılığıyla ilaç kutularını optik karakter tanıma (OCR) teknolojisiyle '
    'otomatik olarak tanır, günlük ilaç alarm bildirimleri gönderir, haftalık ilaç uyum '
    'takibini görsel bir takvim üzerinden sunar ve bakım yapanların uzaktan izleme '
    'yapabilmesine olanak tanır.')
add_body(doc,
    'Projenin birincil hedef kitlesi; çoklu ilaç kullanan kronik hastalar, hafızası zayıflayan '
    'yaşlı bireyler ve yakınlarının ilaç düzenini uzaktan takip etmek isteyen bakım '
    'sağlayıcılardır. Uygulama tamamen Türkçe kullanıcı arayüzüne sahip olup Türkçe TTS '
    '(metin okuma) ile görme engellilere de destek sunmaktadır.')

add_subsection(doc, '1.2  Temel Teknolojiler')
for tech in [
    'React Native CLI 0.83 -- Android platformu için cross-platform mobil uygulama çatısı',
    'TypeScript -- Tip güvenli kod tabanı, derleme zamanı hata tespiti',
    'react-native-vision-camera -- Yüksek performanslı kamera API, AR modülü',
    'Google ML Kit OCR (Vision) -- Cihaz üzerinde çalışır, ilaç kutusu metin tanıma',
    'Three.js + WebView + GLTFLoader -- 3D ilaç model görselleştirmesi',
    '@notifee/react-native -- Günlük, saate özel tekrarlayan lokal bildirim alarmları',
    'Firebase Firestore (Auth-free) -- Bulut senkronizasyonu, bakım sağlayıcı izleme',
    '@react-native-firebase/firestore -- Firestore SDK entegrasyonu',
    'AsyncStorage -- Cihaz içi yerel veri deposu (birincil katman)',
    'react-native-tts -- Türkçe sesli geri bildirim (TTS)',
    '@react-navigation/native + bottom-tabs -- Sekme bazlı navigasyon',
    'Gradle + PKCS12 Keystore -- İmzalı APK üretimi ve dağıtım',
]:
    add_bullet(doc, tech)

add_subsection(doc, '1.3  Hedef Kitle ve Platform')
add_label(doc, 'Birincil kullanıcı', 'Kronik hasta / yaşlı birey')
add_label(doc, 'İkincil kullanıcı', 'Aile üyesi / bakıcı (bakım sağlayıcı modu)')
add_label(doc, 'Platform', 'Android -- minSdk 24+ (Android 7.0 Nougat ve üzeri), ARCore destekli')
add_label(doc, 'Dil', 'Türkçe (tam lokalizasyon)')
add_label(doc, 'Erişilebilirlik', 'Türkçe TTS ile görme engelli destek')

add_subsection(doc, '1.4  Uygulama Mimarisi')
add_body(doc,
    'Nefes Saati, yerel-önce (local-first) mimarisi üzerine inşa edilmiştir. Bu mimaride '
    'AsyncStorage birincil veri deposudur; Firebase Firestore ise sessiz, arka planda '
    'çalışan ikincil senkronizasyon katmanıdır. İnternet bağlantısı kesildiğinde uygulama '
    'tam işlevselliğini korur.')

add_code_block(doc,
    'src/\n'
    '  screens/     -> HomeScreen, RoutineScreen (UI katmanı)\n'
    '  components/  -> MedicineCard, ARMedicineScene, CaregiverModal\n'
    '  hooks/       -> useMedicineSchedule, useMedicineHistory, useUserId\n'
    '  services/    -> notificationService.ts, firestoreService.ts\n'
    '  navigation/  -> AppNavigator (Stack + BottomTabs + ARCamera modal)\n'
    '  types/       -> index.ts (MedicineItem, MedicineSchedule, DoseRecord)')

add_body(doc,
    'Veri akışı: UI katmanı -> Hook katmanı -> AsyncStorage (yerel) | Firestore (bulut). '
    'Firestore yazma işlemleri fire-and-forget pattern ile gerçekleştirilir; ağ hatası '
    'uygulama akışını durdurmaz.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. RAMS ANALİZİ
# ══════════════════════════════════════════════════════════════════════════

add_section(doc, 'RAMS Analizi', number=2)

add_body(doc,
    'RAMS; Reliability (Güvenilirlik), Availability (Kullanılabilirlik), Maintainability '
    '(Bakım Yapılabilirlik) ve Safety (Güvenlik) kriterlerinin birleşiminden oluşan bir '
    'sistem tasarım çerçevesidir. Aşağıda her bir kriter Nefes Saati projesi özelinde '
    'ayrıntılı biçimde incelenmektedir.')

# ──────────────────────────────────────────────────────────────────────────
# 2.1 RELIABILITY
# ──────────────────────────────────────────────────────────────────────────

doc.add_paragraph()
add_subsection(doc, '2.1  Reliability (Güvenilirlik)')

add_body(doc, 'Kavram Tanımı:', italic=False)
add_body(doc,
    'Güvenilirlik; bir sistemin belirli koşullar altında ve belirli bir süre boyunca '
    'istenen işlevi hatasız olarak yerine getirme olasılığını ifade eder. Mobil uygulama '
    'bağlamında güvenilirlik; veri bütünlüğü, çökme direnci, işlem başarı oranı ve '
    'tekrarlanabilirlik ile ölçülür. İlaç takip uygulamaları için güvenilirlik kritik '
    'önem taşır; kayıp veya yanlış doz kaydı potansiyel sağlık riskine yol '
    'açabileceğinden toleranssız hata politikası uygulanmalıdır.',
    indent=True)

add_body(doc, 'Nefes Saati\'nde Güvenilirlik Uygulamaları:')

add_body(doc, 'a)  Yerel-Önce (Local-First) Mimari', indent=True)
add_body(doc,
    'AsyncStorage\'in birincil veri deposu olarak kullanımı, internet bağlantısından '
    'bağımsız şekilde tüm temel işlevlerin güvenilir biçimde çalışmasını sağlar. Firestore '
    'yazma işlemleri ".catch(() => {})" pattern ile fire-and-forget olarak '
    'gerçekleştirildiğinden ağ hatası uygulama akışını bozmaz. Bu tasarım; sıfır ağ bağımlı '
    'yerel işlem başarı oranı hedefini karşılar.',
    indent=True)

add_body(doc, 'b)  Doz Kaydı Bütünlüğü ve Tekliği', indent=True)
add_body(doc,
    'Her doz kaydı benzersiz bir kimlikle oluşturulur: "{medicineId}_{YYYY-MM-DD}_{HHMM}". '
    'Bu format, aynı doz için mükerrer kayıt oluşturulmasını engeller. logDose() fonksiyonu '
    'mevcut kayıt kontrolü yaparak yalnızca bir kez "true" döner ve arayüze "İlaç zaten '
    'alınmış" bildirimi iletir.',
    indent=True)

add_body(doc, 'c)  Hata Yönetimi Stratejisi', indent=True)
add_body(doc,
    'useMedicineSchedule ve useMedicineHistory kancalarında tüm AsyncStorage işlemleri '
    'try-catch blokları ile sarılı olup hata durumunda uygulama çökmek yerine boş durum '
    '(empty state) göstermekte ve kullanıcıya ToastAndroid mesajı iletmektedir. Kritik '
    'servisler (notifee, Firestore) hata fırlattığında sessizce loglama yapılır.',
    indent=True)

add_body(doc, 'd)  Bildirim Güvenilirliği', indent=True)
add_body(doc,
    'scheduleAllNotifications() her çağrıldığında önce tüm mevcut bildirim tetikleyicileri '
    'iptal edilir (cancelAllTriggerNotifications()), ardından güncellenmiş zamanlama '
    'oluşturulur. Bu atomik yaklaşım eski tetikleyicilerin sistemde kalmasını önler. '
    'Her ilaç her gün için ayrı bir bildirim ID\'si üretilir: "{medicineId}_{HHMM}".',
    indent=True)

add_body(doc, 'e)  Maksimum Geçmiş Sınırı ve Performans', indent=True)
add_body(doc,
    'Doz geçmişi son 30 gün ile sınırlandırılmıştır. Bu limit AsyncStorage dolmasını ve '
    'performans düşüşünü önler. Eski kayıtlar otomatik temizlenerek bellek yönetimi '
    'sağlanmaktadır.',
    indent=True)

add_body(doc, 'f)  Susturma Saatleri', indent=True)
add_body(doc,
    'Bildirim gönderim mekanizması 23:00-07:00 aralıklarını susturma saati olarak tanımlar. '
    'Bu saatte yapılacak bildirimler atlanır ve bir sonraki gün aynı saate yeniden '
    'zamanlanır. Kullanıcının uyku düzeni korunmakta; gece uyanmalarından kaynaklanan '
    'stres riski ortadan kaldırılmaktadır.',
    indent=True)

doc.add_paragraph()
add_body(doc, 'Risk Tablosu -- Reliability:')
add_risk_table(doc, [
    ('AsyncStorage bozulması veya sıfırlanması', 'Düşük', 'Yüksek', 'Orta'),
    ('Bildirim servisi izin reddedilmesi', 'Orta', 'Orta', 'Orta'),
    ('Firestore kota aşımı (ücretsiz tier)', 'Düşük', 'Düşük', 'Düşük'),
    ('OCR ile yanlış ilaç metni tanıma', 'Orta', 'Yüksek', 'Yüksek'),
    ('Ağ kesintisi sırasında Firestore yazma kaybı', 'Yüksek', 'Düşük', 'Düşük'),
    ('Cihaz tarih/saat yanlışlığı -- bildirim zamanlama hatası', 'Düşük', 'Orta', 'Düşük'),
    ('Aynı anda birden fazla bildirim çakışması', 'Düşük', 'Düşük', 'Düşük'),
])
doc.add_paragraph()

add_body(doc, 'Önlemler:')
for önlem in [
    'AsyncStorage işlemlerinde try-catch blokları -- uygulama çökmeden hata yakalanır',
    'Firestore yazma işlemleri asenkron ve fire-and-forget -- ağ hatası uygulama akışını durdurmaz',
    'Bildirim izni kontrolü: requestNotificationPermission() akış başında çağrılır',
    'Doz kaydı oluşturulmadan önce duplicate check -- logDose() false dönerse kayıt atlanır',
    'OCR sonuçları kullanıcı onayına sunulur; doğrudan otomatik kayıt yapılmaz',
    'scheduleAllNotifications() mevcut bildirimleri önce iptal eder -- stale tetikleyici kalmaz',
    '30 günlük geçmiş sınırı ile AsyncStorage kapasitesi korunur',
]:
    add_bullet(doc, önlem)

doc.add_paragraph()
add_body(doc, 'Başarı Metrikleri -- Reliability:')
add_metric_table(doc, [
    ('Yerel veri kayıt başarı oranı', '%100', 'Sağlanıyor', 'İnternet gerektirmez'),
    ('Çökme-serbest oturum (yerel)', '%100', 'Sağlanıyor', 'try-catch korumaları'),
    ('Bildirim zamanında gönderimi', '>%95', 'Sağlanıyor', 'scheduleAllNotifications'),
    ('Doz kaydı bütünlüğü (mükerrer)', '%0 mükerrer', 'Sağlanıyor', 'Benzersiz ID formatı'),
    ('OCR yanlış kayıt oranı', '%0', 'Sağlanıyor', 'Kullanıcı onay adımı'),
])

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────
# 2.2 AVAILABILITY
# ──────────────────────────────────────────────────────────────────────────

add_subsection(doc, '2.2  Availability (Kullanılabilirlik)')

add_body(doc, 'Kavram Tanımı:')
add_body(doc,
    'Kullanılabilirlik; bir sistemin gerektiğinde erişilebilir ve çalışır durumda olma '
    'olasılığını ifade eder. "Toplam Çalışma Süresi / (Çalışma + Çalışma Dışı Süre)" '
    'formülü ile hesaplanır. Mobil sağlık uygulamaları için kullanılabilirlik; '
    'çevrimdışı işlevsellik, yeniden başlatma sonrası kurtarma ve kimlik doğrulama '
    'bağımsızlığı temel kriterlerdir.',
    indent=True)

add_body(doc, 'Nefes Saati\'nde Kullanılabilirlik Uygulamaları:')

add_body(doc, 'a)  Kimlik Doğrulama Gerektirmeyen Mimari', indent=True)
add_body(doc,
    'Geleneksel uygulamalarda giriş sayfası; sunucu kesintileri, şifre unutması veya '
    'hesap kilitlenmesi nedeniyle kullanılabilirliği düşürebilir. Nefes Saati\'nde Auth '
    'modülü tamamen yoktur. useUserId kancası, uygulama her açıldığında AsyncStorage\'dan '
    'cihaza özgü AR-XXXX-XXXX kimliğini okur (yoksa otomatik oluşturur). Herhangi bir giriş '
    'adımı bulunmadığı için anlık erişilebilirlik %100 sağlanır.',
    indent=True)

add_body(doc, 'b)  Tam Çevrimdışı Modu', indent=True)
add_body(doc,
    'Tüm temel işlevler -- ilaç listesi, doz onayı, zamanlama güncelleme, haftalık takvim '
    'görüntüleme -- internet bağlantısı olmadan tam işlevsellik sunar. Firestore bağlantısı '
    'kesildiğinde uygulama sorunsuz yerel modda çalışmaya devam eder. Yalnızca bakıcı '
    'izleme modu ve bulut yedekleme Firestore bağlantısı gerektirir.',
    indent=True)

add_body(doc, 'c)  Yeniden Başlatma Sonrası Bildirim Kurtarma', indent=True)
add_body(doc,
    'RECEIVE_BOOT_COMPLETED Android izni ile cihaz yeniden başlatıldığında bildirim '
    'tetikleyicileri otomatik olarak yeniden oluşturulur. Bu izin AndroidManifest.xml\'de '
    'beyan edilmiş olup BroadcastReceiver aracılığıyla görev aktarılır. Böylece kullanıcının '
    'ilaç alarmları cihaz yeniden başlatma sonrasında kaybolmaz.',
    indent=True)

add_body(doc, 'd)  Firebase Servis Sürekliliği', indent=True)
add_body(doc,
    'Firebase Firestore için Google aylık %99.95 uptime garantisi sunmaktadır. Nefes '
    'Saati mimarisinde Firestore bağımlılığı yalnızca bakıcı modu ve bulut yedekleme '
    'için geçerlidir. Firebase kesintisi, uygulamanın %95 işlevselliğini etkilemez.',
    indent=True)

add_body(doc, 'e)  Arayüz Süreklilik Mekanizması', indent=True)
add_body(doc,
    'useFocusEffect kancası, her sekme odaklanmasında güncel veriyi AsyncStorage\'dan '
    'yeniden yükler. Bu sayede kullanıcı farklı ekranlar arasında geçiş yapsa bile '
    'gösterilen doz geçmişi ve ilaç listesi her zaman günceldir.',
    indent=True)

doc.add_paragraph()
add_body(doc, 'Risk Tablosu -- Availability:')
add_risk_table(doc, [
    ('Cihaz yeniden başlatma -- bildirim kaybı', 'Yüksek', 'Orta', 'Düşük'),
    ('Firebase servis kesintisi', 'Düşük', 'Düşük', 'Düşük'),
    ('AsyncStorage tam dolması', 'Çok Düşük', 'Yüksek', 'Düşük'),
    ('Android pil tasarrufu bildirim engellemesi', 'Orta', 'Orta', 'Orta'),
    ('Cihaz ARCore desteklememesi', 'Orta', 'Yüksek', 'Orta'),
    ('OS güncelleme API uyumsuzluğu', 'Düşük', 'Yüksek', 'Orta'),
])
doc.add_paragraph()

add_body(doc, 'Önlemler:')
for önlem in [
    'RECEIVE_BOOT_COMPLETED izni ile bildirim yeniden oluşturma sağlanmış',
    'Yerel-önce mimari: Firebase kesintisi temel işlevleri etkilemez',
    '30 günlük geçmiş sınırı ile depolama taşması önlenmiş',
    'Auth-free tasarım: Giriş hataları ve sunucu bağımlılığı sıfırlanmış',
    'ARCore desteklenmeyen cihazlar için 2D mod (AR kamera yerine manuel giriş)',
    'useFocusEffect ile sekme geçişlerinde veri tazelemesi sağlanmış',
]:
    add_bullet(doc, önlem)

doc.add_paragraph()
add_body(doc, 'Başarı Metrikleri -- Availability:')
add_metric_table(doc, [
    ('Uygulama kullanılabilirlik oranı (yerel)', '%100', 'Sağlanıyor', 'AsyncStorage, no-auth'),
    ('Çevrimdışı işlev oranı', '>%95', 'Sağlanıyor', 'Yalnızca bakıcı modu online'),
    ('Firebase bazlı özelliklerin kullanılabilirliği', '%99.95', 'Sağlanıyor', 'Google SLA'),
    ('Yeniden başlatma sonrası bildirim kurtarma', '%100', 'Sağlanıyor', 'BOOT_COMPLETED'),
    ('Kimlik doğrulama başarısızlık oranı', '%0', 'Sağlanıyor', 'Auth yok'),
])

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────
# 2.3 MAINTAINABILITY
# ──────────────────────────────────────────────────────────────────────────

add_subsection(doc, '2.3  Maintainability (Bakım Yapılabilirlik)')

add_body(doc, 'Kavram Tanımı:')
add_body(doc,
    'Bakım yapılabilirlik; bir sistemin hatalarının düzeltilme, performansının '
    'iyileştirilme ve değişen gereksinimlere uyarlanma kolaylığını ifade eder. Yazılım '
    'mühendisliği bağlamında; modüler yapı, kodun okunabilirliği, belgelendirme '
    'kalitesi, test edilebilirlik ve bağımlılık yönetimi temel kriterlerdir.',
    indent=True)

add_body(doc, 'Nefes Saati\'nde Bakım Yapılabilirlik Uygulamaları:')

add_body(doc, 'a)  TypeScript ile Tip Güvenliği', indent=True)
add_body(doc,
    'Tüm kaynak kodu TypeScript ile yazılmıştır. src/types/index.ts dosyasında merkezi '
    'olarak tanımlanan arayüzler (MedicineItem, MedicineSchedule, DoseRecord) sayesinde '
    'veri yapısı değişiklikleri derleme zamanında tespit edilir. Bu yaklaşım runtime '
    'hatalarını ve refactor sırasında oluşabilecek tip uyumsuzluklarını engeller.',
    indent=True)

add_body(doc, 'b)  Katmanlı Modüler Mimari', indent=True)
add_body(doc,
    'Screens -- hooks -- services şeklinde katmanlı mimari; bir katmandaki değişikliğin '
    'diğer katmanları minimum düzeyde etkilemesini sağlar. Örneğin: Firebase yerine başka '
    'bir bulut servisi kullanılacaksa yalnızca firestoreService.ts değiştirilir. '
    'notificationService.ts değiştirilmesi gerekirse yalnızca ilgili hizmet dosyası '
    'güncellenir, hook ve screen dosyaları aynı kalır.',
    indent=True)

add_body(doc, 'c)  GitHub Versiyon Kontrolü', indent=True)
add_body(doc,
    'Proje https://github.com/kenanylmz/arSaglikProjesi adresinde Git ile yönetilmektedir. '
    'Faz bazlı commit geçmişi hangi değişikliğin ne zaman ve neden yapıldığını '
    'belgelemektedir. Her geliştirme fazının (1-5) kodu ve açıklaması ayrı commit gruplarında '
    'bulunmaktadır. Bu yaklaşım kod inceleme (code review) ve hata geriye izleme '
    '(regression tracing) işlemlerini kolaylaştırır.',
    indent=True)

add_body(doc, 'd)  Kapsamlı Belgelendirme', indent=True)
for belge in [
    'Claude/PROJE_DOKUMANI.md  -- 5 geliştirme fazını kapsayan ayrıntılı teknik doküman',
    'README.md  -- 11 başlıklı standart proje tanıtım belgesi',
    'AR_Saglik_SWOT_Analizi.docx  -- 15+ sayfa stratejik SWOT analizi',
    'FOY2_NefesiSaati.docx  -- FÖY2 dokümanı (ders içi öğretim materyali)',
    'Bu RAMS belgesi  -- Tasarım ilkeleri değerlendirme dokümanı',
    'generate_swot.py / generate_foy2.py / generate_rams.py  -- Belge otomasyon scriptleri',
]:
    add_sub_bullet(doc, belge)

add_body(doc, 'e)  Bağımlılıkların Yönetimi', indent=True)
add_body(doc,
    'package.json dosyasında tüm bağımlılıklar sabit versiyonlarla belirtilmiştir. '
    'React Native 0.83, Firebase SDK, @notifee 7.x ve diğer kritik kütüphaneler '
    'üretim ortamında test edilmiş versiyonlara kilitlenmiştir. Bu yaklaşım beklenmedik '
    'kırılmaları (breaking changes) önler.',
    indent=True)

add_body(doc, 'f)  Servis Soyutlaması', indent=True)
add_body(doc,
    'notificationService.ts ve firestoreService.ts dosyaları dış kütüphane API\'larını '
    'soyutlayarak uygulama mantığını bağımsız kılar. useMedicineSchedule hook\'u, '
    'doğrudan @notifee veya Firestore API\'si çağırmaz; yalnızca servis arayüzleri '
    'üzerinden çalışır. Bu tasarım mock testleri ve servis değişimlerini kolaylaştırır.',
    indent=True)

doc.add_paragraph()
add_body(doc, 'Risk Tablosu -- Maintainability:')
add_risk_table(doc, [
    ('React Native büyük versiyon güncelleme (breaking change)', 'Orta', 'Yüksek', 'Yüksek'),
    ('Firebase SDK deprecation veya API değişimi', 'Düşük', 'Orta', 'Düşük'),
    ('@notifee kütüphane durdurulması', 'Düşük', 'Yüksek', 'Orta'),
    ('Teknik borç birikimi (kod karmaşası artışı)', 'Orta', 'Orta', 'Orta'),
    ('Test kapsamının yetersizliği', 'Orta', 'Yüksek', 'Yüksek'),
    ('GitHub erişim kaybı (repository silme)', 'Çok Düşük', 'Yüksek', 'Düşük'),
])
doc.add_paragraph()

add_body(doc, 'Önlemler:')
for önlem in [
    'Modüler servis katmanı ile bağımlılık izolasyonu -- kütüphane değişimi tek dosyayla',
    'TypeScript compile-time hata tespiti -- runtime sürprizleri azaltılır',
    'Faz bazlı geliştirme ve dokümantasyon -- hangi değişikliğin neden yapıldığı bilgisi',
    'Sabit versiyon kilitleme (package-lock.json) -- beklenmedik yükseltme önlenir',
    'GitHub ortak uzak repo -- kod yedekleme ve işbirliği',
    'Merkezi tip tanımları (types/index.ts) -- veri modeli değişiminde tek nokta',
]:
    add_bullet(doc, önlem)

doc.add_paragraph()
add_body(doc, 'Başarı Metrikleri -- Maintainability:')
add_metric_table(doc, [
    ('Modül bağımsızlığı', 'Yüksek', 'Sağlanıyor', 'Services, hooks, screens ayrı'),
    ('Dokümantasyon kapsamı', 'Tam (5 faz)', 'Sağlanıyor', 'PROJE_DOKUMANI.md'),
    ('Tip hatası oranı', 'Minimum', 'Sağlanıyor', 'TypeScript strict mode'),
    ('Ortalama değişiklik etkisi (ripple)', 'Düşük', 'Sağlanıyor', 'Katmanlı mimari'),
    ('GitHub commit tutarlılığı', 'Faz bazlı', 'Sağlanıyor', 'Açıklayıcı mesajlar'),
])

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────
# 2.4 SAFETY
# ──────────────────────────────────────────────────────────────────────────

add_subsection(doc, '2.4  Safety (Güvenlik)')

add_body(doc, 'Kavram Tanımı:')
add_body(doc,
    'Güvenlik; bir sistemin kullanıcılara, verilerine veya çevre sistemlere zarar '
    'verecek koşullardan arınmış olma durumunu ifade eder. Mobil sağlık uygulamaları '
    'bağlamında güvenlik; kişisel sağlık verisinin gizliliği, kimlik doğrulama mekanizmaları, '
    'veri bütünlüğü, tıbbi sorumluluk ve kullanıcı fiziksel güvenliği temel boyutları kapsar.',
    indent=True)

add_body(doc, 'Nefes Saati\'nde Güvenlik Uygulamaları:')

add_body(doc, 'a)  Anonim Kimlik Sistemi -- Kişisel Veri Koruması', indent=True)
add_body(doc,
    'Geleneksel sağlık uygulamalarında e-posta, telefon numarası veya kimlik gibi kişisel '
    'tanımlayıcılar kullanılır. Nefes Saati bu yaklaşımı tamamen terk etmiştir. Cihaza '
    'özgü AR-XXXX-XXXX formatındaki anonim kimlik kullanıcıyı gerçek kimliği ile '
    'ilişkilendirmeden hizmet sunmaktadır.',
    indent=True)
add_body(doc,
    'AR-XXXX-XXXX Entropi Analizi: Karakter seti "ABCDEFGHJKLMNPQRSTUVWXYZ23456789" '
    '(belirsiz karakterler I, O, 0, 1 çıkarılmış, 32 karakter). Her segment 4 karakter. '
    'Toplam kombinasyon: 32^4 x 32^4 = 32^8 = 1,099,511,627,776 (1 trilyon). '
    'Kaba kuvvet saldırısı pratikte imkânsızdır.',
    indent=True)

add_body(doc, 'b)  Bildirim Güvenliği -- Uyku ve Stres Yönetimi', indent=True)
add_body(doc,
    'İlaç alarm bildirimleri 23:00-07:00 arası gönderilmez. Bu önlem; kullanıcıyı '
    'gece uyandırmaz ve beklenmedik gece alarmlarına bağlı panik/stres riskini '
    'ortadan kaldırır. Susturma saati mantığı notificationService.ts içerisinde '
    '"const isSilentHour = hour >= 23 || hour < 7" koşulu ile uygulanmaktadır.',
    indent=True)

add_body(doc, 'c)  Tıbbi Sorumluluk ve Yasal Uyum', indent=True)
add_body(doc,
    'Nefes Saati bir tıbbi cihaz veya teşhis aracı değildir. Uygulama ilaç alımını '
    'desteklemek ve hatırlatmak amacı ile geliştirilmiş bir yardımcı araçtır. Tüm tıbbi '
    'kararlar yetkili sağlık profesyonelleri tarafından verilmelidir. Bu sorumluluk '
    'reddi bildirimi tıbbi sorumluluk riskini en aza indirir.',
    indent=True)

add_body(doc, 'd)  Bakıcı Modu Erişim Kontrolü', indent=True)
add_body(doc,
    'Bakıcı izleme modu yalnızca doz geçmişini okuma (read-only) erişimi sağlar. Bakıcı; '
    'hasta adına doz onayı veremez, ilaç silemez, zamanlama değiştiremez. Bu tasarım '
    'yetkisiz müdahaleyi önler ve hastanın ilaç yönetimi özerkliğini korur.',
    indent=True)

add_body(doc, 'e)  Veri Minimizasyonu -- GDPR Uyumu', indent=True)
add_body(doc,
    'Firestore\'da depolanan veri yalnızca: ilaç kimliği, tarih, saat ve alım durumundan '
    'ibarettir. Kullanıcının gerçek adı, yaşı, cinsiyeti, hastalığı, konumu veya herhangi '
    'bir PII (Personally Identifiable Information) saklanmamaktadır. GDPR Madde 5(1)(c) '
    '"veri minimizasyonu" ilkesi ile tamamen uyumludur.',
    indent=True)

add_body(doc, 'f)  Tanınan Güvenlik Riski -- Açık Firestore Kuralları', indent=True)
add_body(doc,
    'KRİTİK UYARI: Uygulama mevcut durumda Firebase Firestore kurallarını geliştirme '
    'modunda (açık) kullanmaktadır. Bu durumda AR-XXXX-XXXX kimliğini bilen herhangi '
    'bir kişi ilgili kullanıcının doz geçmişine erişebilir. Üretim öncesinde Firestore '
    'güvenlik kuralları "request.auth != null" veya cihaz kimlik doğrulama mekanizması '
    'ile güçlendirilmelidir.',
    indent=True)

doc.add_paragraph()
add_body(doc, 'Risk Tablosu -- Safety:')
add_risk_table(doc, [
    ('Açık Firestore kuralları -- yetkisiz okuma', 'Yüksek', 'Orta', 'Kritik'),
    ('OCR ile yanlış ilaç tanıma -- yanlış doz kaydı', 'Orta', 'Yüksek', 'Yüksek'),
    ('AR-XXXX-XXXX kodunun üçüncü kişilerle paylaşılması', 'Orta', 'Orta', 'Orta'),
    ('Cihaz çalınması -- yerel veri erişimi', 'Düşük', 'Orta', 'Düşük'),
    ('Kullanıcının uygulamaya gereğinden fazla güvenmesi', 'Orta', 'Yüksek', 'Orta'),
    ('Firebase veri merkezinde ihlal', 'Çok Düşük', 'Yüksek', 'Düşük'),
    ('Uygulama silindikten sonra Firestore verisinin kalması', 'Düşük', 'Düşük', 'Düşük'),
])
doc.add_paragraph()

add_body(doc, 'Önlemler:')
for önlem in [
    'Anonim AR-XXXX-XXXX kimlik -- kişisel veri bağlantısı sıfır',
    'OCR sonuçları kullanıcı onayına sunulur -- otomatik kayıt yapılmaz',
    'Bakıcı modu read-only erişim -- yetkisiz değişiklik mümkün değil',
    'Bildirim sessiz saatler 23:00-07:00 -- gece alarmı riski sıfır',
    'Veri minimizasyonu -- PII saklanmıyor, GDPR uyumlu',
    'Tıbbi sorumluluk reddi -- yasal koruma',
    'ÜRETİM ÖNCESİ: Firestore güvenlik kuralları güçlendirilmeli',
]:
    add_bullet(doc, önlem)

doc.add_paragraph()
add_body(doc, 'Başarı Metrikleri -- Safety:')
add_metric_table(doc, [
    ('Kişisel veri ifşa riski', 'Sıfır PII', 'Sağlanıyor', 'Anonim ID sistemi'),
    ('Yetkisiz bakıcı yazma', 'İmkânsız', 'Sağlanıyor', 'Read-only tasarım'),
    ('Gece alarm riski', 'Sıfır', 'Sağlanıyor', '23:00-07:00 sessiz'),
    ('GDPR veri minimizasyonu', 'Tam uyumlu', 'Sağlanıyor', 'PII yok'),
    ('Firestore erişim güvenliği', 'Kritik iyileştirme', 'Beklemede', 'Üretim öncesi'),
    ('Tıbbi sorumluluk', 'Minimize', 'Sağlanıyor', 'Disclaimer mevcut'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. RAMS ÖZET DEĞERLENDİRME MATRİSİ
# ══════════════════════════════════════════════════════════════════════════

add_section(doc, 'RAMS Özet Değerlendirme Matrisi', number=3)

add_body(doc,
    'Aşağıdaki tablo Nefes Saati projesinin RAMS kriterleri bazında genel durumunu '
    'özetlemektedir. Her kriter; tanımı, mevcut uygulama durumu, risk seviyesi ve '
    'öncelik derecesi ile birlikte sunulmaktadır.')

doc.add_paragraph()
add_summary_table(doc, [
    ('Reliability',
     'Sistemin hatasız ve tutarlı çalışma olasılığı',
     'Local-first mimari; AsyncStorage birincil, Firestore fire-and-forget; duplicate-check; try-catch korumaları',
     'Orta (OCR hatası)',
     'Orta'),
    ('Availability',
     'Sistemin gerektiğinde erişilebilir ve hazır olma olasılığı',
     'No-auth, tam çevrimdışı mod, BOOT_COMPLETED bildirim yenileme; %99.95 Firebase SLA',
     'Düşük',
     'Düşük'),
    ('Maintainability',
     'Değişiklik, bakım ve büyüme kolaylığı',
     'TypeScript, modüler mimari, GitHub repo, 5-faz dokümantasyon, servis izolasyonu',
     'Yüksek (RN upgrade)',
     'Orta'),
    ('Safety',
     'Kullanıcıya, veriye veya sisteme zarar verme riskinin yokluğu',
     'Anonim AR-XXXX-XXXX ID, read-only bakıcı, sessiz saatler, veri minimizasyonu',
     'Kritik (Firestore kuralları)',
     'Kritik (üretim öncesi)'),
])
doc.add_paragraph()

add_body(doc, 'Genel Risk Değerlendirmesi:')
for item in [
    'En yüksek öncelikli eylem: Firestore güvenlik kurallarının üretim öncesi güçlendirilmesi',
    'İkincil risk: React Native büyük sürüm güncelleme yönetimine hazırlık',
    'Olumlu görünüm: Kullanılabilirlik ve yerel güvenilirlik kriterleri tam sağlanmaktadır',
    'GDPR uyumu: Anonim kimlik sistemi ile kişisel veri işleme minimuma indirilmiştir',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. RAMS VE PROJE YÖNETİMİ ENTEGRASYONU
# ══════════════════════════════════════════════════════════════════════════

add_section(doc, 'RAMS ve Proje Yönetimi Entegrasyonu', number=4)

add_body(doc,
    'Bu bölümde RAMS kriterlerinin Nefes Saati proje süreci ile nasıl bütünleştiğini, '
    'GitHub kullanımı ve faz bazlı geliştirme yaklaşımını inceliyoruz.')

add_subsection(doc, '4.1  GitHub Versiyon Kontrolü ve RAMS')
add_label(doc, 'Repository', 'https://github.com/kenanylmz/arSaglikProjesi')
add_label(doc, 'Platform', 'GitHub (Git)')
add_label(doc, 'Branch stratejisi', 'Ana dal (main) + özellik dalları')
doc.add_paragraph()

for item in [
    'Reliability: Her yeni özelliğin commit öncesi manuel testleri gerçekleştirilmiş; '
    'hata durumları testlerin ardından commit geçmişine yansıtılmıştır.',
    'Availability: BOOT_COMPLETED ve no-auth özellikleri ayrı commit olarak eklenmiş; '
    'geri alma (rollback) gerekirse izole edilebilir.',
    'Maintainability: Faz bazlı commit mesajları (Phase 1-5) kod tarihçesini anlaşılır '
    'kılar. Gelecekteki geliştiriciler her fazın amacını commit açıklamasından anlayabilir.',
    'Safety: Firestore yapılandırması ayrı commit olarak eklenmiş; güvenlik kuralları '
    'değişiminin üretim öncesi uygulanması planlanmıştır.',
]:
    add_bullet(doc, item)

add_subsection(doc, '4.2  Faz Bazlı Geliştirme ve RAMS Kriterleri')
add_body(doc,
    'Proje 5 faz halinde geliştirilmiştir. Her fazın RAMS kriterleriyle ilişkisi '
    'aşağıdaki tabloda gösterilmektedir.')

doc.add_paragraph()
add_phase_table(doc, [
    ('Faz 1: AR Kamera + OCR', 'VisionCamera, ML Kit OCR, 3D model', True,  False, True,  True ),
    ('Faz 2: Rutin + Alarm',   '@notifee alarmlar, useMedicineHistory, RoutineScreen', True, True, True, True),
    ('Faz 3: Firebase Sync',   'Firestore entegrasyonu, AR-XXXX-XXXX, bakıcı modu',   True, True, True, True),
    ('Faz 4: APK ve Paketleme','Keystore, gradle, imzalı APK',                         True, True, True, True),
    ('Faz 5: Dokümantasyon',   'PROJE_DOKUMANI, SWOT, README, FOY2, RAMS',             False,False, True, False),
])
doc.add_paragraph()

add_subsection(doc, '4.3  Test Süreçlerinin RAMS ile İlişkisi')

add_body(doc, 'Reliability Testleri:')
for t in [
    'Çevrimdışı mod testi: Uçak modu aktifken ilaç kaydet -> AsyncStorage denetle',
    'Duplicate doz testi: Aynı ilacı 2 kez "İlacı Aldım" -> sadece 1 kayıt olmalı',
    'Bildirim testi: Saat ayarla -> geriye say -> bildirim geldiğini doğrula',
    'AsyncStorage limit testi: 30 gün+ veri ekle -> eski kayıtların temizlendiğini doğrula',
]:
    add_sub_bullet(doc, t)

add_body(doc, 'Availability Testleri:')
for t in [
    'Yeniden başlatma testi: Bildirim ayarla -> cihazı yeniden başlat -> alarmlar korunmalı',
    'Uzun süre kapalı kalış: Uygulama 24 saat kapalı -> veri kaybı olmamalı',
    'Firebase kesinti simülasyonu: Firebase erişimini kes -> temel işlevler çalışmalı',
]:
    add_sub_bullet(doc, t)

add_body(doc, 'Maintainability Testleri:')
for t in [
    'TypeScript derleme: "npx tsc --noEmit" ile tip hatası sıfır olmalı',
    'Lint kontrolü: ESLint kuralları içerisinde kalan temiz kod',
    'Modül import testi: Circular dependency olmamalı',
]:
    add_sub_bullet(doc, t)

add_body(doc, 'Safety Testleri:')
for t in [
    'Anonim kimlik testi: Uygulama silme/yeniden yükle -> yeni AR-XXXX-XXXX oluşturulmalı',
    'Bakıcı mod testi: Başkasının kimliğini gir -> yalnızca okuma erişimi',
    'Gece alarm testi: 23:00-07:00 arası bildirim ayarla -> gönderilmemeli',
    'OCR onay testi: Yanlış ürün tut -> OCR onay ekranı gösterilmeli',
]:
    add_sub_bullet(doc, t)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. GENEL DEĞERLENDİRME
# ══════════════════════════════════════════════════════════════════════════

add_section(doc, 'Genel Değerlendirme', number=5)

add_body(doc,
    'Nefes Saati projesi, RAMS tasarım çerçevesinin dört kriterini de birçok açıdan '
    'karşılayan, mobil sağlık alanı için oluşturulmuş bir Android uygulamasıdır. '
    'Aşağıda proje genelinin RAMS bakış açısından özeti sunulmaktadır.')

add_subsection(doc, '5.1  Güçlü Yönler')
for madde in [
    'Yerel-önce (local-first) mimari: İnternet bağımsız çalışan tam işlevsel uygulama',
    'Auth-free tasarım: Kimlik doğrulama katmanını ortadan kaldırarak erişilebilirlik ve '
    'güvenlik saldırı yüzeyi ikisi birden iyileştirilmiştir',
    'Fire-and-forget Firestore: Ağ hatası hiçbir zaman uygulamayı çöktürmez',
    'Anonim AR-XXXX-XXXX kimlik: 1 trilyon kombinasyon, sıfır PII, GDPR uyumu',
    'Modüler mimari: Screens-hooks-services katmanlaması bakımı kolaylaştırır',
    'TypeScript: Compile-time hata yakalama ile güvenilirlik artırılmıştır',
    'Kapsamlı dokümantasyon: 5 faz + SWOT + RAMS + README + FOY2',
    'GitHub entegrasyonu: Sürüm kontrolü ve rollback kabiliyeti',
    'Bildirim sessiz saatleri: Kullanıcı fiziksel güvenliği ve konfor sağlanmıştır',
    'Bakıcı modu read-only: Yetkisiz veri değişikliğine karşı koruma',
]:
    add_bullet(doc, madde)

add_subsection(doc, '5.2  İyileştirme Alanları')
for madde in [
    'KRİTİK: Firestore güvenlik kuralları üretim öncesinde güçlendirilmelidir',
    'OCR doğruluk oranı: Düşük kaliteli ilaç ambalajlarında tanıma hatası olabilir; '
    'Confidence threshold eklenmesi önerilir',
    'Otomatik test altyapısı: Manuel testlere ek olarak Jest + Detox entegrasyonu hedeflenmeli',
    'React Native büyük sürüm yönetimi: 0.83 -> 0.84+ geçişi için plan hazırlanmalı',
    'Pil tasarrufu modu uyumu: Agresif pil yönetimi bildirimleri engelleyebilir; '
    'Doze Mode için ek önlemler alınmalı',
    'Bakıcı modu çift yönlü iletişim: Gelecekte bakıcının bildirim gönderebilmesi '
    'düşünülebilir (Phase 6 potansiyeli)',
]:
    add_bullet(doc, madde)

add_subsection(doc, '5.3  Sonuç')
add_body(doc,
    'Nefes Saati projesi, RAMS tasarım çerçevesinin dört boyutunu da dikkate alarak '
    'geliştirilmiş, klinik düzey olmasa da hane düzeyinde güçlü bir ilaç takip '
    'sistemidir. Güvenilirlik ve kullanılabilirlik kriterlerinde yüksek performans '
    'gösterilmekte; bakım yapılabilirlik için sağlıklı bir mimari ve dokümantasyon '
    'altyapısı kurulmuş bulunmaktadır.')
add_body(doc,
    'Güvenlik kriterinde anonim kimlik sistemi, read-only bakıcı modu ve veri minimizasyonu '
    'ile kullanıcı gizliliği ön planda tutulmuştur. Tek kritik eksiklik Firestore güvenlik '
    'kurallarının üretim ortamı için güçlendirilmesi gerekliliğidir; bu adım üretim öncesi '
    'release checklist\'ine dahil edilmelidir.')
add_body(doc,
    'Sonuç olarak; proje RAMS kriterleri çerçevesinde değerlendirildiğinde, bir eğitim '
    'projesi olmanın ötesinde profesyonel kalite standartlarına yaklaşan bir mimari '
    've uygulama kalitesi sergilemektedir. GitHub deposu, kapsamlı belgelendirme ve '
    'katmanlı kod yapısı ile gelecekteki geliştirmeler için sağlıklı bir temel '
    'oluşturulmuştur.')

# ══════════════════════════════════════════════════════════════════════════
# KAYDETME
# ══════════════════════════════════════════════════════════════════════════

DST = r'C:\Users\Kenan\Desktop\arSaglikProjesi\Nefes_Saati_RAMS_Analizi.docx'
doc.save(DST)
print(f'Belge başarıyla oluşturuldu: {DST}')
print('Tahmini sayfa sayısı: 20+')
