"""
AR Sağlık Projesi - SWOT Analizi Word Belgesi Oluşturucu
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Renkler ──────────────────────────────────────────────
COLOR_PRIMARY    = RGBColor(0x0D, 0x6E, 0x6E)  # Teal
COLOR_STRENGTH   = RGBColor(0x27, 0xAE, 0x60)  # Yeşil
COLOR_WEAKNESS   = RGBColor(0xE7, 0x4C, 0x3C)  # Kırmızı
COLOR_OPP        = RGBColor(0x29, 0x80, 0xB9)  # Mavi
COLOR_THREAT     = RGBColor(0xE6, 0x7E, 0x22)  # Turuncu
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK       = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_GRAY       = RGBColor(0x4A, 0x4A, 0x6A)
COLOR_LIGHT_BG   = RGBColor(0xF5, 0xF7, 0xFA)

# ── Yardımcı fonksiyonlar ────────────────────────────────
def set_cell_shading(cell, hex_color):
    """Hücre arka plan rengini ayarla."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_cell_text(cell, text, bold=False, color=COLOR_DARK, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Hücreye formatlı metin ekle."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"


def add_bullet_cell(cell, items, color=COLOR_DARK, bg_hex="FFFFFF"):
    """Hücreye madde işaretli liste ekle."""
    set_cell_shading(cell, bg_hex)
    cell.text = ""
    for i, item in enumerate(items):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.space_before = Pt(3)
        p.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)

        bullet = p.add_run("● ")
        bullet.font.size = Pt(10)
        bullet.font.color.rgb = color
        bullet.font.name = "Calibri"
        bullet.bold = True

        # Başlık ve açıklama ayır
        if ":" in item:
            title, desc = item.split(":", 1)
            title_run = p.add_run(title + ":")
            title_run.font.size = Pt(10.5)
            title_run.font.color.rgb = COLOR_DARK
            title_run.font.name = "Calibri"
            title_run.bold = True

            desc_run = p.add_run(desc)
            desc_run.font.size = Pt(10.5)
            desc_run.font.color.rgb = COLOR_GRAY
            desc_run.font.name = "Calibri"
        else:
            text_run = p.add_run(item)
            text_run.font.size = Pt(10.5)
            text_run.font.color.rgb = COLOR_DARK
            text_run.font.name = "Calibri"


def add_heading_styled(doc, text, level=1, color=COLOR_PRIMARY):
    """Renkli başlık ekle."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return heading


def add_paragraph_styled(doc, text, bold=False, color=COLOR_DARK, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Formatlı paragraf ekle."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_divider(doc):
    """Yatay çizgi ekle."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)


# ── SWOT Verileri ────────────────────────────────────────

STRENGTHS = [
    "OCR Tabanlı İlaç Tanıma (ML Kit): VisionCamera + Google ML Kit metin tanıma entegrasyonu sayesinde hastalar ilaç kutusunu kameraya tutarak saniyeler içinde tanıma yaptırabilir. OCR yaklaşımı, marker bağımlılığını ortadan kaldırarak farklı baskı serisi ve kutu kondisyonlarında çok daha güvenilir sonuçlar verir.",
    "Sesli Komut Sistemi (TTS): Türkçe dil desteğiyle entegre edilen metin-konuşma teknolojisi, görme güçlüğü çeken yaşlı hastalar için kritik bir erişilebilirlik özelliğidir. Hasta ilacını kameraya tuttuğunda 'Bu Sandimmun Neoral. Bir sonraki kullanım saati akşam 8. 3 saat 22 dakika kaldı.' gibi net ve anlaşılır sesli bildirimler alır.",
    "Haftalık Rutinim & İlaç Geçmiş Takibi: 7 günlük grid görünümüyle alınan/kaçırılan/bekleyen her doz anlık izlenebilir. Haftalık uyum yüzdesi otomatik hesaplanır. Bu özellik, hastaların kendi tedavilerine sahip çıkmasını ve bakıcıların uzaktan takibini güçlendirir.",
    "Onaylı Doz Kaydı — 'İlacı Aldım' Sistemi: AR overlay açılınca hasta 'İlacı Aldım' butonuna basarak dozu onaylar. Kayıt AsyncStorage ve Firebase'e eş zamanlı yazılır. Böylece hem çevrimdışı güvenlik hem de bulut senkronizasyonu sağlanır; kaçırılan dozlar geçmişte görünür kalır.",
    "Akıllı Bildirim & Alarm Sistemi: @notifee/react-native entegrasyonu ile her ilaç saati için günlük tekrarlayan Android bildirimleri otomatik planlanır. Kullanıcı saati güncellediğinde alarmlar anında yeniden kurulur. Sessiz saatler (23:00–07:00) korumalıdır.",
    "Firebase Firestore Bulut Senkronizasyonu: Uygulama 'local-first' mimarisiyle çalışır; AsyncStorage birincil, Firestore ikincil depolama katmanıdır. Ağ kesintilerinde uygulama kesintisiz çalışır; bağlantı gelince otomatik senkronizasyon başlar. Cihaz değişikliğinde veri kaybı yaşanmaz.",
    "Ebeveyn / Bakıcı Uzaktan Takip Modu: Hasta yakını, Rutinim ekranındaki bakıcı butonuna tıklayıp hastanın kısa ID'sini (AR-XXXX-XXXX formatı) girerek o günün tüm doz durumunu anlık görebilir. Firebase üzerinden çalışır; iki farklı cihaz arasında gerçek zamanlı takip mümkündür.",
    "Auth-Free Cihaz Kimliği Mimarisi: Kayıt veya giriş gerektirmez. Her cihaza otomatik üretilen 8 karakterlik okunabilir ID (ör. AR-K3M8-PQ7X) kalıcı kimlik olarak kullanılır. Yaşlı hastaların şifre ve hesap yönetimiyle uğraşmaması için bilinçli alınmış kritik bir tasarım kararıdır.",
    "Kişiselleştirilebilir İlaç Saatleri ile Tam Entegrasyon: Saat düzenlendiğinde AsyncStorage, Firebase Firestore ve @notifee alarmları tek seferde güncellenir. Üç sistem arasındaki tutarlılık otomatik sağlanır; veri tutarsızlığı riski yoktur.",
    "Clean Code Mimarisi & Local-First Tasarım: constants, hooks, services, components, screens, navigation katmanlarıyla sürdürülebilir TypeScript mimarisi. Firestore hataları sessizce yakalanır; uygulama asla Firestore'a bağımlı kalmaz. Offline-first yaklaşım, yaşlı hasta cihazlarının bağlantı kalitesi göz önüne alındığında kritik önem taşır.",
]

WEAKNESSES = [
    "Sınırlı İlaç Veritabanı: Uygulama şu an yalnızca iki ilaç (Sandimmun Neoral ve Deltacortril) tanıyabilmektedir. Karaciğer nakli sonrası hastaların kullanabileceği Tacrolimus (Prograf), Mycophenolate Mofetil (CellCept), Ursodiol gibi ek ilaçlar henüz OCR anahtar kelime listesine eklenmemiştir.",
    "OCR Tanıma Güvenilirliği: ML Kit metin tanıma, ilaç kutusunun fiziksel durumuna (yıpranmış, buruşmuş, farklı dil baskısı) ve ortam aydınlatmasına bağlı olarak hatalı sonuç verebilir. Kutu üzerindeki anahtar kelime eşleştirme yöntemi, görüntü sınıflandırma modellerine kıyasla daha kırılgandır.",
    "Açık Firestore Güvenlik Kuralları: Mevcut güvenlik kuralları 'allow read, write: if true' şeklinde yapılandırılmıştır. Bu durum, URL'i bilen herhangi birinin koleksiyona erişebileceği anlamına gelir. Klinik ortama taşınmadan önce kural sıkılaştırması zorunludur.",
    "Tek Platform Desteği (Android): Şu anki sürüm yalnızca Android cihazları desteklemektedir. iOS kullanan hasta ve bakıcıların uygulamadan yararlanamaması, özellikle bakıcı takip modunun kullanımını kısıtlar.",
    "Doktor-Hasta Entegrasyonu Yok: Doktor veya eczacının hastanın ilaç programını uzaktan görüntülemesi veya güncellemesi mümkün değildir. Bakıcı modu yalnızca aile üyelerini kapsamakta; klinisyen entegrasyonu bulunmamaktadır.",
    "İlaç Etkileşimi Bilgisi Yok: Uygulama ilaçlar arası etkileşim, yan etkiler veya kontraendikasyon bilgisi sunmamaktadır. Ciclosporin ile greyfurt etkileşimi gibi kritik besin-ilaç uyarıları henüz eklenmemiştir.",
    "Erişilebilirlik Sınırlamaları: Ekran okuyucu (TalkBack) uyumluluğu, renk körlüğü desteği ve motor engelli kullanıcılar için büyük dokunma alanları gibi gelişmiş erişilebilirlik özellikleri henüz implemente edilmemiştir.",
    "Bakıcı Modu — Tek Yönlü Veri Akışı: Bakıcı şu an yalnızca izleyebilir; hastanın saat programını uzaktan değiştirme, mesaj gönderme veya acil uyarı tetikleme yetkisi yoktur. Ebeveyn kontrolü ileride çift yönlü yapıya dönüştürülmelidir.",
    "Firebase Bağımlılığı ve Tek Tedarikçi Riski: Google'ın Firebase fiyatlandırma politikası veya servis koşullarında yapacağı değişiklikler, özellikle veri hacmi arttıkça maliyet riski oluşturabilir. Alternatif backend mimarisine geçiş planı henüz yoktur.",
    "İnternet Bağımsız Güncelleme Eksikliği: Yeni ilaç ekleme veya OCR anahtar kelime genişletme için uygulama güncellemesi gerekir. Uzaktan ilaç listesi güncelleme mekanizması bulunmamakta; doktor reçete değişikliklerinde gecikme yaşanabilir.",
]

OPPORTUNITIES = [
    "Sağlık Sektöründe Dijital Dönüşüm: Küresel dijital sağlık pazarı 2025-2030 arasında yıllık %25+ büyüme öngörülmektedir. Nefes Saati, AR + bulut senkronizasyonu + bakıcı takibi kombinasyonuyla bu trendin erken dönem referans uygulamalarından biri olarak konumlanabilir.",
    "Hastane Taburculuk Süreci Entegrasyonu: Karaciğer nakli merkezleri ve organ nakli koordinasyon birimleriyle işbirliği yapılarak uygulama taburculuk paketinin standart parçası haline getirilebilir. Bakıcı modu, hastane ile aile arasındaki iletişim köprüsünü dijitalleştirir.",
    "Firebase Altyapısının Genişletilmesi: Mevcut Firestore yapısı; ilaç etkileşimi kontrolü, doktor portalı, anlık bildirim (bakıcıya 'doz kaçırıldı' uyarısı) ve çok kullanıcılı aile grubu yönetimi gibi özellikler için sağlam bir temel oluşturmaktadır.",
    "Yapay Zeka Destekli Tanıma: Mevcut OCR altyapısı, TensorFlow Lite veya on-device ML modeli ile güçlendirilerek tablet/kapsül görüntüsünden tanıma yapılabilir hale getirilebilir. Bu, kutunun açık ya da yıpranmış olduğu durumlarda güvenilirliği önemli ölçüde artırır.",
    "Çoklu Hastalık Alanlarına Genişleme: Aynı mimari; böbrek nakli, kalp nakli, diyabet, onkoloji gibi alanlara adapte edilebilir. Her hastalık için özelleştirilmiş OCR anahtar kelime listesi ve ilaç protokolü eklenebilir. Firebase, çok tenant mimarisine uygundur.",
    "Bakıcı Modunun Klinisyen Portala Dönüştürülmesi: Bakıcı ID sorgulama altyapısı, doktor ve hemşire için web paneli (React veya Next.js) ile genişletilebilir. Doktor birden fazla hastanın uyum raporunu tek ekrandan görebilir.",
    "Akademik Araştırma ve Klinik Çalışmalar: Firestore'da biriken gerçek zamanlı ilaç uyum verileri; 'AR + bulut tabanlı hatırlatma sistemlerinin organ nakli sonrası uyuma etkisi' konusunda özgün akademik yayın ve tez materyali sunmaktadır.",
    "Devlet Sağlık Politikalarıyla Uyum: Türkiye'nin dijital sağlık dönüşümü stratejisi ve e-Nabız entegrasyonlarıyla uyumlu konumlanarak SGK veya Sağlık Bakanlığı desteği alınabilir. Bakıcı takip modu, yaşlı bakımı politikalarına doğrudan katkı sağlar.",
    "IoT ve Akıllı İlaç Kutusu Entegrasyonu: Bluetooth destekli akıllı ilaç kutuları ile entegrasyon sağlanarak fiziksel ilaç alımının otomatik Firebase'e yazılması ve bakıcıya anlık bildirim gönderilmesi mümkün hale getirilebilir.",
    "Giyilebilir Cihaz & Bildirim Genişlemesi: @notifee altyapısı Wear OS uyumlu akıllı saatlere genişletilebilir. Hasta bileğinde titreşimli ilaç hatırlatması, özellikle yaşlı kullanıcılar için telefon bildiriminden daha güvenilir bir alternatif sunar.",
    "Eczane ve İlaç Firması Ortaklıkları: Novartis ve Pfizer gibi üreticilerle sponsorluk anlaşmaları yapılarak ilaç veritabanı genişletilebilir, 3D model kütüphanesi zenginleştirilebilir. Bakıcı izleme verileri anonim olarak ilaç firmaları için değerli uyum araştırması sağlar.",
]

THREATS = [
    "Açık Firestore Kurallarından Kaynaklanan Güvenlik Riski: Geliştirme aşamasında kullanılan 'allow read, write: if true' kuralları, gerçek hasta verilerinin kötü niyetli kişilerce okunmasına veya değiştirilmesine olanak tanır. Klinik kullanım öncesi bu kuralların kesinlikle sıkılaştırılması gerekir.",
    "Yasal ve Düzenleyici Riskler: Tıbbi cihaz yazılımı olarak sınıflandırılma riski mevcuttur. T.C. İlaç ve Tıbbi Cihaz Kurumu (TİTCK) onayı gerekebilir. Firestore'da sağlık verisi saklanması, KVKK ve GDPR kapsamında ek yükümlülükler doğurur.",
    "Hasta Güvenliği Riskleri: Hatalı OCR tanıma veya yanlış saat bilgisi, hastanın yanlış ilacı almasına ya da dozu kaçırmasına neden olabilir. Karaciğer nakli hastalarında bu durum organ reddi gibi hayati sonuçlara yol açabilir. 'İlacı Aldım' onayı bu riski azaltsa da ortadan kaldırmaz.",
    "Firebase Bağımlılığı ve Tek Tedarikçi Riski: Google Firebase'in fiyatlandırma politikası değiştiğinde veya servis kesintisi yaşandığında bakıcı modu ve bulut senkronizasyonu devre dışı kalır. Uygulama offline çalışmaya devam etse de veri paylaşımı etkilenir.",
    "Rekabet Ortamı: Medisafe, MyTherapy ve Pill Reminder gibi olgun uygulamalar milyonlarca kullanıcıya sahiptir. Bu uygulamalar bakıcı bildirimi, çok ilaçlı takip ve doktor entegrasyonu gibi özelliklere zaten sahipken, Nefes Saati'nin AR ve OCR farkı tek başına yeterli diferansiyasyon sağlamayabilir.",
    "Yaşlı Kullanıcı Adaptasyon Zorluğu: Hedef kitle olan 60+ yaş grubu, AR kamera kullanımı, ID paylaşımı ve bildirim yönetimine adapte olmakta güçlük çekebilir. Uygulamanın görece zenginleşen özellik seti, başlangıçtaki minimal tasarım ilkesiyle çelişme riski taşır.",
    "Yanlış Tıbbi Bilgi Riski: İlaç bilgilerinin güncel ve doğru tutulması sorumluluğu kritiktir. Güncel olmayan prospektüs bilgileri veya doz saatleri hukuki sorumluluk doğurabilir. Firebase'de saklanan verinin doğruluğu uygulama tarafından doğrulanamamaktadır.",
    "Siber Güvenlik Tehditleri: Sağlık verilerini hedefleyen siber saldırılar artmaktadır. AsyncStorage ve Firestore'daki verilerin şifrelenmemiş olması, cihaz çalınması veya kötü amaçlı yazılım bulaşması durumunda hasta bilgilerini riske atar.",
    "Platform ve Kütüphane Eskime Riski: React Native 0.83, @notifee, @react-native-firebase gibi bağımlılıkların büyük versiyon güncellemeleri kırıcı değişiklikler getirebilir. Özellikle yeni React Native mimarisi (Fabric/TurboModules) bazı kütüphanelerin yeniden yazılmasını gerektirebilir.",
    "Cihaz Donanım Gereksinimleri: Kamera sürekli kullanımı, Firestore ağ bağlantısı ve @notifee alarm işlemleri pil tüketimini artırır. Yaşlı hastaların genellikle eski ve düşük kapasiteli cihaz kullanması, uzun süreli kullanım deneyimini olumsuz etkileyebilir.",
]

# ── AR Teknolojisi Genel SWOT ────────────────────────────

AR_STRENGTHS = [
    "Gerçek Dünya ile Dijital Bilgi Birleşimi: AR, fiziksel nesnelerin üzerine dijital içerik bindirerek kullanıcıya bağlamsal ve anlık bilgi sunar. Bu, sağlık alanında ilaç tanıma, cerrahi navigasyon ve hasta eğitiminde devrim niteliğindedir.",
    "Sezgisel Etkileşim Modeli: Kullanıcılar doğal hareketlerle (bakma, tutma, yönlendirme) dijital içerikle etkileşim kurar. Bu, ekran tabanlı arayüzlere kıyasla öğrenme eğrisini önemli ölçüde düşürür.",
    "Mekansal Anlayış: AR teknolojisi ortamı 3D olarak algılar ve anlam çıkarır. SLAM (Simultaneous Localization and Mapping) algoritmaları sayesinde gerçek dünyadaki yüzeyleri, nesneleri ve derinliği anlayabilir.",
    "Çoklu Duyusal Deneyim: Görsel (3D model), işitsel (sesli komut) ve haptik (titreşim) geri bildirimleri birleştirerek çok kanallı bir kullanıcı deneyimi sunar.",
    "Eğitim ve Simülasyon Gücü: Tıp öğrencileri anatomik yapıları AR ile inceleyebilir, cerrahlar ameliyat öncesi 3D planlama yapabilir. AR, teorik bilgiyi pratik uygulamaya dönüştürme potansiyeli taşır.",
]

AR_WEAKNESSES = [
    "Donanım Bağımlılığı: AR deneyimi, cihazın kamerası, işlemcisi, sensörleri ve ekran kalitesine doğrudan bağlıdır. Düşük seviye cihazlarda performans sorunları ve deneyim kalitesi düşüşü kaçınılmazdır.",
    "Pil Tüketimi: Sürekli kamera, GPS, sensör ve GPU kullanımı cihaz pilini hızla tüketir. Mobil AR uygulamalarında ortalama %300-400 daha fazla pil tüketimi gözlemlenir.",
    "Çevresel Koşullara Duyarlılık: Yetersiz aydınlatma, düz yüzey eksikliği, reflektif materyaller ve hızlı hareket, AR deneyiminin kalitesini ciddi şekilde etkiler.",
    "Kullanıcı Yorgunluğu: Uzun süreli AR kullanımı göz yorgunluğu, baş dönmesi ve hareket tutarsızlığı (motion sickness) gibi fizyolojik sorunlara yol açabilir.",
    "Gizlilik Endişeleri: AR uygulamaları sürekli kamera erişimi gerektirir. Bu durum, kullanıcıların ve çevredekilerin gizliliği konusunda ciddi etik ve hukuki sorular doğurur.",
]

AR_OPPORTUNITIES = [
    "5G ve Edge Computing: 5G ağlarının yaygınlaşması, bulut tabanlı AR işlemeyi mümkün kılarak cihaz bağımlılığını azaltacak ve daha karmaşık AR deneyimlerini mobil cihazlarda gerçekleştirmeyi sağlayacaktır.",
    "AR Gözlüklerin Yaygınlaşması: Apple Vision Pro, Meta Quest, ve diğer AR gözlüklerin tüketici pazarına girmesiyle, hands-free AR deneyimleri günlük hayatın parçası haline gelecektir.",
    "WebAR Standartları: Tarayıcı tabanlı AR (WebXR API) olgunlaştıkça, uygulama indirme gerekliliği ortadan kalkacak ve AR erişimi demokratikleşecektir.",
    "Yapay Zeka + AR Sinerjisi: Makine öğrenmesi modelleriyle güçlendirilen AR, nesne tanıma, sahne anlama ve tahminsel içerik yerleştirme konularında çığır açıcı gelişmeler sunacaktır.",
    "Endüstri 4.0 Entegrasyonu: Üretim, lojistik, bakım-onarım ve sağlık sektörlerinde AR tabanlı iş akışları standart hale gelecek, verimlilik artışları %20-35 arasında öngörülmektedir.",
]

AR_THREATS = [
    "Platform Parçalanması: ARCore (Google), ARKit (Apple), OpenXR ve çeşitli SDK'lar arasındaki uyumsuzluklar, geliştiricilerin çoklu platform desteği sağlamasını zorlaştırır ve maliyetleri artırır.",
    "Düzenleyici Belirsizlik: AR içeriklerinin denetlenmesi, sorumluluk paylaşımı ve dijital içeriklerle fiziksel dünyayı birleştirmenin hukuki çerçevesi henüz net değildir.",
    "Sosyal Kabul Bariyerleri: Kamuya açık alanlarda AR cihaz kullanımı, toplumsal tepki ve gizlilik endişeleri yaratabilir. 'Google Glass sendromu' hâlâ sektör için bir uyarıdır.",
    "Güvenlik Açıkları: AR overlay'leri manipüle edilerek kullanıcıları yanıltma (AR spoofing/phishing) potansiyeli, yeni siber güvenlik tehdit vektörleri oluşturmaktadır.",
    "Yüksek Geliştirme Maliyetleri: Kaliteli AR içerik üretimi (3D modelleme, mekansal ses, gerçek zamanlı rendering) geleneksel mobil uygulama geliştirmeden 3-5 kat daha maliyetlidir.",
]


# ── Belge Oluşturma ──────────────────────────────────────

def create_document():
    doc = Document()

    # Sayfa ayarları
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── KAPAK SAYFASI ────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    # Başlık
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NEFES SAATİ")
    run.font.size = Pt(36)
    run.font.color.rgb = COLOR_PRIMARY
    run.font.name = "Calibri"
    run.bold = True

    # Alt başlık
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SWOT ANALİZİ")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_GRAY
    run.font.name = "Calibri"

    # Açıklama
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Karaciğer Nakli Sonrası Hasta İlaç Takip Sistemi\nArtırılmış Gerçeklik (AR) Tabanlı Mobil Uygulama")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRAY
    run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Tarih ve bilgiler
    info_items = [
        "Teknoloji: React Native CLI 0.83 + VisionCamera + ML Kit + Three.js + Firebase Firestore + @notifee",
        "Platform: Android",
        "İlaçlar: Sandimmun Neoral (Ciclosporin) | Deltacortril (Prednisolone)",
        "Tarih: Nisan 2026",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRAY
        run.font.name = "Calibri"

    # Sayfa sonu
    doc.add_page_break()

    # ── İÇİNDEKİLER ─────────────────────────────────────
    add_heading_styled(doc, "İÇİNDEKİLER", level=1)
    doc.add_paragraph()

    toc_items = [
        ("1.", "Giriş ve Proje Tanımı"),
        ("2.", "SWOT Analizi Nedir?"),
        ("3.", "Proje SWOT Analizi"),
        ("   3.1.", "Güçlü Yönler (Strengths)"),
        ("   3.2.", "Zayıf Yönler (Weaknesses)"),
        ("   3.3.", "Fırsatlar (Opportunities)"),
        ("   3.4.", "Tehditler (Threats)"),
        ("4.", "SWOT Matrisi"),
        ("5.", "AR Teknolojisi Genel SWOT Analizi"),
        ("   5.1.", "AR Güçlü Yönler"),
        ("   5.2.", "AR Zayıf Yönler"),
        ("   5.3.", "AR Fırsatlar"),
        ("   5.4.", "AR Tehditler"),
        ("6.", "Stratejik Öneriler"),
        ("7.", "Sonuç ve Değerlendirme"),
    ]

    for num, title_text in toc_items:
        p = doc.add_paragraph()
        run_num = p.add_run(num + " ")
        run_num.font.size = Pt(12)
        run_num.font.color.rgb = COLOR_PRIMARY
        run_num.font.name = "Calibri"
        run_num.bold = True

        run_title = p.add_run(title_text)
        run_title.font.size = Pt(12)
        run_title.font.color.rgb = COLOR_DARK
        run_title.font.name = "Calibri"

    doc.add_page_break()

    # ── 1. GİRİŞ ────────────────────────────────────────
    add_heading_styled(doc, "1. Giriş ve Proje Tanımı", level=1)

    intro_text = (
        "Nefes Saati, karaciğer nakli sonrası tedavi sürecindeki yaşlı hastaların ilaç yönetimini "
        "kolaylaştırmak amacıyla geliştirilmiş, artırılmış gerçeklik (AR) tabanlı Android mobil "
        "uygulamasıdır. React Native CLI 0.83 altyapısı üzerine inşa edilen uygulama; VisionCamera, "
        "Google ML Kit OCR, Three.js 3D model görüntüleyici, Firebase Firestore ve @notifee bildirim "
        "sistemi gibi modern teknolojileri tek çatı altında birleştirmektedir."
    )
    add_paragraph_styled(doc, intro_text)

    intro2 = (
        "Karaciğer nakli sonrası hastalar, bağışıklık sistemini baskılamak ve organ reddini önlemek için "
        "ömür boyu ilaç kullanmak zorundadır. Bu ilaçların düzenli ve doğru zamanda alınması hayati "
        "önem taşır. Özellikle yaşlı hasta grubunda ilaç uyumsuzluğu (non-adherence) oranları %40-60 "
        "arasında seyretmekte; bu durum organ reddi, hastaneye yeniden yatış ve ölüm riskini "
        "önemli ölçüde artırmaktadır."
    )
    add_paragraph_styled(doc, intro2)

    intro3 = (
        "Uygulama; ilaç kutusu OCR tanıma, gerçek zamanlı 3D model ve Türkçe sesli bildirim özelliklerinin "
        "yanı sıra, 'İlacı Aldım' onay sistemi, haftalık rutinim ekranı, Firebase tabanlı bulut senkronizasyonu "
        "ve ebeveyn/bakıcı uzaktan takip modunu da kapsamaktadır. Tüm bu özellikler, hasta ve hasta yakınlarını "
        "tedavi sürecinde aktif birer katılımcıya dönüştürmeyi hedeflemektedir."
    )
    add_paragraph_styled(doc, intro3)

    # İlaç bilgisi tablosu
    doc.add_paragraph()
    add_paragraph_styled(doc, "Projede Kullanılan İlaçlar:", bold=True, color=COLOR_PRIMARY, size=13)

    med_table = doc.add_table(rows=3, cols=5)
    med_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Başlıklar
    headers = ["İlaç Adı", "Etken Madde", "Doz", "Üretici", "Kullanım"]
    for i, h in enumerate(headers):
        add_cell_text(med_table.rows[0].cells[i], h, bold=True, color=COLOR_WHITE, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(med_table.rows[0].cells[i], "0D6E6E")

    # Neoral
    neoral_data = ["Sandimmun Neoral", "Ciclosporin", "100 mg", "Novartis", "Günde 2 kez\n(08:00 - 20:00)"]
    for i, d in enumerate(neoral_data):
        add_cell_text(med_table.rows[1].cells[i], d, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(med_table.rows[1].cells[i], "FFF8E7")

    # Deltacortril
    delta_data = ["Deltacortril", "Prednisolone", "5 mg", "Pfizer", "Günde 1 kez\n(10:00)"]
    for i, d in enumerate(delta_data):
        add_cell_text(med_table.rows[2].cells[i], d, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(med_table.rows[2].cells[i], "E8EEF5")

    doc.add_page_break()

    # ── 2. SWOT ANALİZİ NEDİR? ──────────────────────────
    add_heading_styled(doc, "2. SWOT Analizi Nedir?", level=1)

    swot_desc = (
        "SWOT Analizi, bir projenin, organizasyonun veya ürünün stratejik konumunu değerlendirmek için "
        "kullanılan dört boyutlu bir analiz çerçevesidir. SWOT kısaltması İngilizce Strengths (Güçlü Yönler), "
        "Weaknesses (Zayıf Yönler), Opportunities (Fırsatlar) ve Threats (Tehditler) kelimelerinin "
        "baş harflerinden oluşur."
    )
    add_paragraph_styled(doc, swot_desc)

    doc.add_paragraph()

    # SWOT açıklama tablosu
    swot_info = doc.add_table(rows=2, cols=2)
    swot_info.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells_info = [
        ("GÜÇLÜ YÖNLER (S)", "İç faktörler: Projenin sahip olduğu avantajlar, benzersiz özellikler ve rekabet üstünlükleri.", "27AE60", "E8F8F5"),
        ("ZAYIF YÖNLER (W)", "İç faktörler: Projenin eksiklikleri, iyileştirmeye açık alanlar ve dezavantajlar.", "E74C3C", "FDEDEC"),
        ("FIRSATLAR (O)", "Dış faktörler: Projenin büyümesi ve gelişmesi için yararlanılabilecek dış koşullar.", "2980B9", "EBF5FB"),
        ("TEHDİTLER (T)", "Dış faktörler: Projenin başarısını riske atan dış tehlikeler ve engeller.", "E67E22", "FDF2E9"),
    ]

    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for (row, col), (title_text, desc_text, header_color, bg_color) in zip(positions, cells_info):
        cell = swot_info.rows[row].cells[col]
        cell.text = ""

        # Başlık
        p_title = cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.space_before = Pt(8)
        run = p_title.add_run(title_text)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(int(header_color[:2], 16), int(header_color[2:4], 16), int(header_color[4:], 16))
        run.font.name = "Calibri"
        run.bold = True

        # Açıklama
        p_desc = cell.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.space_after = Pt(8)
        run = p_desc.add_run(desc_text)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GRAY
        run.font.name = "Calibri"

        set_cell_shading(cell, bg_color)

    doc.add_page_break()

    # ── 3. PROJE SWOT ANALİZİ ───────────────────────────
    add_heading_styled(doc, "3. Proje SWOT Analizi", level=1)
    add_paragraph_styled(
        doc,
        "Bu bölümde AR Sağlık Projesi'nin güçlü yönleri, zayıf yönleri, fırsatları ve tehditleri "
        "ayrıntılı olarak incelenmektedir.",
        color=COLOR_GRAY
    )
    doc.add_paragraph()

    # 3.1 Güçlü Yönler
    add_heading_styled(doc, "3.1 Güçlü Yönler (Strengths)", level=2, color=COLOR_STRENGTH)
    for i, item in enumerate(STRENGTHS, 1):
        title_part, desc_part = item.split(":", 1) if ":" in item else (f"Madde {i}", item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        num_run = p.add_run(f"  {i}. ")
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = COLOR_STRENGTH
        num_run.font.name = "Calibri"
        num_run.bold = True

        title_run = p.add_run(title_part.strip() + ":")
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLOR_DARK
        title_run.font.name = "Calibri"
        title_run.bold = True

        desc_run = p.add_run(desc_part.strip())
        desc_run.font.size = Pt(11)
        desc_run.font.color.rgb = COLOR_GRAY
        desc_run.font.name = "Calibri"

    doc.add_page_break()

    # 3.2 Zayıf Yönler
    add_heading_styled(doc, "3.2 Zayıf Yönler (Weaknesses)", level=2, color=COLOR_WEAKNESS)
    for i, item in enumerate(WEAKNESSES, 1):
        title_part, desc_part = item.split(":", 1) if ":" in item else (f"Madde {i}", item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        num_run = p.add_run(f"  {i}. ")
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = COLOR_WEAKNESS
        num_run.font.name = "Calibri"
        num_run.bold = True

        title_run = p.add_run(title_part.strip() + ":")
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLOR_DARK
        title_run.font.name = "Calibri"
        title_run.bold = True

        desc_run = p.add_run(desc_part.strip())
        desc_run.font.size = Pt(11)
        desc_run.font.color.rgb = COLOR_GRAY
        desc_run.font.name = "Calibri"

    doc.add_page_break()

    # 3.3 Fırsatlar
    add_heading_styled(doc, "3.3 Fırsatlar (Opportunities)", level=2, color=COLOR_OPP)
    for i, item in enumerate(OPPORTUNITIES, 1):
        title_part, desc_part = item.split(":", 1) if ":" in item else (f"Madde {i}", item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        num_run = p.add_run(f"  {i}. ")
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = COLOR_OPP
        num_run.font.name = "Calibri"
        num_run.bold = True

        title_run = p.add_run(title_part.strip() + ":")
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLOR_DARK
        title_run.font.name = "Calibri"
        title_run.bold = True

        desc_run = p.add_run(desc_part.strip())
        desc_run.font.size = Pt(11)
        desc_run.font.color.rgb = COLOR_GRAY
        desc_run.font.name = "Calibri"

    doc.add_page_break()

    # 3.4 Tehditler
    add_heading_styled(doc, "3.4 Tehditler (Threats)", level=2, color=COLOR_THREAT)
    for i, item in enumerate(THREATS, 1):
        title_part, desc_part = item.split(":", 1) if ":" in item else (f"Madde {i}", item)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        num_run = p.add_run(f"  {i}. ")
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = COLOR_THREAT
        num_run.font.name = "Calibri"
        num_run.bold = True

        title_run = p.add_run(title_part.strip() + ":")
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = COLOR_DARK
        title_run.font.name = "Calibri"
        title_run.bold = True

        desc_run = p.add_run(desc_part.strip())
        desc_run.font.size = Pt(11)
        desc_run.font.color.rgb = COLOR_GRAY
        desc_run.font.name = "Calibri"

    doc.add_page_break()

    # ── 4. SWOT MATRİSİ ─────────────────────────────────
    add_heading_styled(doc, "4. SWOT Matrisi (Özet Tablo)", level=1)
    add_paragraph_styled(
        doc,
        "Aşağıdaki matris, projenin SWOT analizini özet olarak sunmaktadır.",
        color=COLOR_GRAY
    )
    doc.add_paragraph()

    # 2x2 matris tablosu
    matrix = doc.add_table(rows=3, cols=3)
    matrix.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Sol üst köşe boş
    add_cell_text(matrix.rows[0].cells[0], "", size=10)
    set_cell_shading(matrix.rows[0].cells[0], "0D6E6E")

    # Üst başlıklar
    add_cell_text(matrix.rows[0].cells[1], "OLUMLU\n(Destekleyici)", bold=True, color=COLOR_WHITE, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(matrix.rows[0].cells[1], "0D6E6E")
    add_cell_text(matrix.rows[0].cells[2], "OLUMSUZ\n(Engelleyici)", bold=True, color=COLOR_WHITE, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(matrix.rows[0].cells[2], "0D6E6E")

    # İç faktörler satırı
    add_cell_text(matrix.rows[1].cells[0], "İÇ\nFAKTÖRLER", bold=True, color=COLOR_WHITE, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(matrix.rows[1].cells[0], "0D6E6E")

    strength_summary = [s.split(":")[0] for s in STRENGTHS[:6]]
    add_bullet_cell(matrix.rows[1].cells[1], strength_summary, COLOR_STRENGTH, "E8F8F5")

    weakness_summary = [w.split(":")[0] for w in WEAKNESSES[:6]]
    add_bullet_cell(matrix.rows[1].cells[2], weakness_summary, COLOR_WEAKNESS, "FDEDEC")

    # Dış faktörler satırı
    add_cell_text(matrix.rows[2].cells[0], "DIŞ\nFAKTÖRLER", bold=True, color=COLOR_WHITE, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(matrix.rows[2].cells[0], "0D6E6E")

    opp_summary = [o.split(":")[0] for o in OPPORTUNITIES[:6]]
    add_bullet_cell(matrix.rows[2].cells[1], opp_summary, COLOR_OPP, "EBF5FB")

    threat_summary = [t.split(":")[0] for t in THREATS[:6]]
    add_bullet_cell(matrix.rows[2].cells[2], threat_summary, COLOR_THREAT, "FDF2E9")

    doc.add_page_break()

    # ── 5. AR TEKNOLOJİSİ GENEL SWOT ────────────────────
    add_heading_styled(doc, "5. Artırılmış Gerçeklik (AR) Teknolojisi - Genel SWOT Analizi", level=1)
    add_paragraph_styled(
        doc,
        "Bu bölümde, projenin temelini oluşturan Artırılmış Gerçeklik teknolojisinin genel olarak "
        "güçlü ve zayıf yönleri, sektörel fırsatları ve tehditleri değerlendirilmektedir.",
        color=COLOR_GRAY
    )
    doc.add_paragraph()

    # AR SWOT - her bölüm
    ar_sections = [
        ("5.1 AR - Güçlü Yönler", AR_STRENGTHS, COLOR_STRENGTH),
        ("5.2 AR - Zayıf Yönler", AR_WEAKNESSES, COLOR_WEAKNESS),
        ("5.3 AR - Fırsatlar", AR_OPPORTUNITIES, COLOR_OPP),
        ("5.4 AR - Tehditler", AR_THREATS, COLOR_THREAT),
    ]

    for section_title, items, color in ar_sections:
        add_heading_styled(doc, section_title, level=2, color=color)
        for i, item in enumerate(items, 1):
            title_part, desc_part = item.split(":", 1) if ":" in item else (f"Madde {i}", item)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)

            num_run = p.add_run(f"  {i}. ")
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = color
            num_run.font.name = "Calibri"
            num_run.bold = True

            title_run = p.add_run(title_part.strip() + ":")
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = COLOR_DARK
            title_run.font.name = "Calibri"
            title_run.bold = True

            desc_run = p.add_run(" " + desc_part.strip())
            desc_run.font.size = Pt(11)
            desc_run.font.color.rgb = COLOR_GRAY
            desc_run.font.name = "Calibri"

        doc.add_paragraph()

    doc.add_page_break()

    # ── 6. STRATEJİK ÖNERİLER ───────────────────────────
    add_heading_styled(doc, "6. Stratejik Öneriler", level=1)

    strategies = [
        ("SO Stratejileri (Güçlü Yönler + Fırsatlar)", COLOR_STRENGTH, [
            "Bakıcı modu + Firebase altyapısını kullanarak hastane taburculuk paketine entegre olun. Hasta uyum verilerinin Firestore'dan klinik panele aktarılması, hastane yöneticileri için güçlü bir satış argümanıdır.",
            "OCR + 3D model + TTS altyapısını ML tabanlı görüntü sınıflandırmasıyla güçlendirin. On-device TensorFlow Lite modeli, kutu durumundan bağımsız çok daha güvenilir tanıma sağlar.",
            "React Native'in çapraz platform avantajını kullanarak iOS sürümünü çıkarın. Bakıcı modunun iOS'ta da çalışması, aile üyelerinin farklı platformlardan takip yapmasını sağlar.",
            "Akademik yayınlar için Firestore'daki anonim uyum verilerini kullanın. 'Bulut senkronizasyonlu AR hatırlatma sistemlerinin organ nakli sonrası uyuma etkisi' tezi için gerçek veri mevcuttur.",
            "Offline-first mimarisini öne çıkararak kırsal bölgelerdeki hastanelere ulaşın; bağlantı gelince senkronize olan yapı, düşük internet kalitesinde bile güvenilirlik sağlar.",
        ]),
        ("WO Stratejileri (Zayıf Yönler + Fırsatlar)", COLOR_OPP, [
            "Firestore güvenlik kurallarını Firebase Authentication veya cihaz ID doğrulamasıyla sıkılaştırın. Bu adım, gerçek hasta verisiyle çalışmak için ön koşuldur.",
            "İlaç veritabanı sınırlılığını eczane ve ilaç firması ortaklıklarıyla aşın; Tacrolimus, CellCept gibi nakil ilaçlarını OCR listesine ve Firestore'a ekleyin.",
            "Bakıcı modunu çift yönlü hale getirin: bakıcı yalnızca izlemekle kalmayıp hastaya 'ilacını al' bildirimi gönderebilsin, saat programını uzaktan düzenleyebilsin.",
            "Doktor-hasta entegrasyonu için Firestore üzerine bir web paneli (React/Next.js) inşa edin. Mevcut veri modeli bu genişlemeye hazırdır, ek backend mimarisi gerekmez.",
            "Erişilebilirlik sınırlamalarını TalkBack uyumluluğu ve sesli navigasyon ile aşarak engelli hasta grubuna hizmet verin; yaşlı dostu tasarım ilkesi bu entegrasyonu kolaylaştırmaktadır.",
        ]),
        ("ST Stratejileri (Güçlü Yönler + Tehditler)", COLOR_THREAT, [
            "Firestore güvenlik riskine karşı, 'İlacı Aldım' onay sisteminin oluşturduğu audit trail'i veri bütünlüğü kanıtı olarak belgelendirin. KVKK uyumu için kamera verilerinin cihaz üzerinde işlendiğini ve sunucuya gönderilmediğini açıkça ilan edin.",
            "Hasta güvenliği için OCR tanıma sonucuna 'Bu ilacı mı kullanıyorsunuz? Eğer yanlışsa taramayı iptal edin.' onay adımı ekleyin. 'İlacı Aldım' butonu bu zinciri tamamlar.",
            "Rekabet tehdidine karşı, Auth-free cihaz ID + bakıcı takip + AR tanıma kombinasyonunu farklılaştırıcı özellik olarak öne çıkarın; patent veya faydalı model başvurusu yapın.",
            "Firebase tek tedarikçi riskine karşı, servis katmanını soyutlayın (firestoreService.ts zaten bunu yapmaktadır). Gerektiğinde Supabase veya MongoDB Atlas'a geçiş minimum eforla mümkün olsun.",
            "Yasal riskleri minimize etmek için 'tıbbi bilgi aracı' konumlamasını koruyun, 'tıbbi cihaz' sınıflandırmasından kaçının; uygulama içi sorumluluk reddi beyanları ekleyin.",
        ]),
        ("WT Stratejileri (Zayıf Yönler + Tehditler)", COLOR_WEAKNESS, [
            "Açık Firestore kuralları + siber güvenlik tehdidi: Üretim öncesinde mutlaka kural sıkılaştırması yapın. Geçiş dönemi için en azından cihaz ID eşleşmesi zorunluluğu getiren kural seti hazırlayın.",
            "Sınırlı ilaç veritabanı + yanlış tanıma riski: Tanınmayan ilaçlar için 'Bu ilacı tanıyamadım, lütfen doktorunuza danışın' mesajı zorunlu tutun; kullanıcının manuel ilaç seçimi yapabileceği fallback mod ekleyin.",
            "Yaşlı adaptasyon zorluğu + özellik artışı çelişkisi: Yeni özellikleri (bakıcı modu, rutinim ekranı) 'ileri kullanıcı' segmentine yönlendirin; ana ekranı sade tutarak temel kullanım akışını koruyun.",
            "Firebase bağımlılığı + platform eskime riski: Kritik iş mantığını (doz hesaplama, alarm planlama) servis katmanında izole tutun. Her kütüphane güncellemesinde entegrasyon testleri yapın.",
            "Düzenleyici belirsizliğe karşı Sağlık Bakanlığı ve TİTCK ile erken aşamada iletişime geçin; Firestore'da saklanan sağlık verisinin hukuki statüsünü netleştirmek klinik kullanım öncesinde zorunludur.",
        ]),
    ]

    for title_text, color, items in strategies:
        add_heading_styled(doc, title_text, level=2, color=color)
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.5)

            num_run = p.add_run(f"{i}. ")
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = color
            num_run.font.name = "Calibri"
            num_run.bold = True

            text_run = p.add_run(item)
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = COLOR_GRAY
            text_run.font.name = "Calibri"

    doc.add_page_break()

    # ── 7. SONUÇ ─────────────────────────────────────────
    add_heading_styled(doc, "7. Sonuç ve Değerlendirme", level=1)

    conclusion_paras = [
        "Nefes Saati, karaciğer nakli sonrası ilaç yönetimi alanında artırılmış gerçeklik, bulut "
        "senkronizasyonu ve bakıcı takip sistemi özelliklerini bir arada sunan kapsamlı bir dijital "
        "sağlık uygulamasıdır. SWOT analizi sonuçları, projenin güçlü bir teknolojik ve mimari temele "
        "sahip olduğunu; ancak güvenlik, veritabanı genişliği ve platform çeşitliliği konularında "
        "iyileştirme fırsatları barındırdığını ortaya koymaktadır.",

        "Projenin en belirgin güçlü yönleri; OCR tabanlı ilaç tanıma, Türkçe sesli bildirim, "
        "'İlacı Aldım' onay sistemi, haftalık rutinim takibi ve ebeveyn/bakıcı uzaktan izleme modudur. "
        "Bu özellikler kombinasyonu, uygulamayı mevcut ilaç hatırlatma uygulamalarından net biçimde "
        "farklılaştırmaktadır. Özellikle local-first + Firebase senkronizasyonu mimarisi, hem "
        "çevrimdışı güvenilirliği hem de bulut tabanlı paylaşımı aynı anda garanti altına alır.",

        "Açık Firestore güvenlik kuralları, sınırlı ilaç veritabanı ve tek platform desteği gibi "
        "zayıf yönler, uygulamanın klinik ortama taşınabilmesi için öncelikli ele alınması gereken "
        "konulardır. Bu zayıflıkların giderilmesi, projenin gerçek hasta gruplarıyla kullanılabilmesi "
        "için zorunlu ön koşullardır.",

        "Fırsatlar açısından bakıldığında, Firebase altyapısının doktor portalına, IoT entegrasyonuna "
        "ve akademik veri analizine dönüştürülme potansiyeli oldukça yüksektir. Mevcut Firestore veri "
        "modeli, bu genişlemeler için yeniden yapılandırma gerektirmeksizin temel oluşturmaktadır. "
        "Ancak bu fırsatların değerlendirilmesi; güvenlik kuralı sıkılaştırması, yasal düzenlemeler "
        "ve hasta güvenliği riskleriyle paralel yürütülmelidir.",

        "Stratejik yol haritası olarak; kısa vadede Firestore güvenlik kurallarının sıkılaştırılması "
        "ve OCR veritabanının genişletilmesi, orta vadede iOS desteği ve çift yönlü bakıcı modu, "
        "uzun vadede doktor portalı ve ML tabanlı görüntü tanıma sistemi hedeflenmelidir. "
        "Bu aşamalı yaklaşım, projenin sürdürülebilir büyümesini ve klinik etkisini maksimize edecektir.",

        "Sonuç olarak Nefes Saati; doğru stratejik kararlarla, karaciğer nakli sonrası ilaç "
        "uyumluluğu sorununa yenilikçi, güvenilir ve kapsamlı bir çözüm sunma potansiyeline sahip, "
        "umut vadeden bir dijital sağlık girişimidir.",
    ]

    for para_text in conclusion_paras:
        add_paragraph_styled(doc, para_text)

    # ── ALT BİLGİ ────────────────────────────────────────
    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AR Sağlık Projesi - SWOT Analizi Raporu\nNisan 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_GRAY
    run.font.name = "Calibri"
    run.italic = True

    return doc


# ── Ana çalıştırma ───────────────────────────────────────
if __name__ == "__main__":
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "AR_Saglik_SWOT_Analizi.docx")

    print("SWOT Analizi Word belgesi oluşturuluyor...")
    doc = create_document()
    doc.save(output_path)
    print(f"Belge başarıyla oluşturuldu: {output_path}")
    print(f"Toplam sayfa: ~15+")
    print(f"Proje SWOT: {len(STRENGTHS)} güçlü yön, {len(WEAKNESSES)} zayıf yön, {len(OPPORTUNITIES)} fırsat, {len(THREATS)} tehdit")
    print(f"AR SWOT: {len(AR_STRENGTHS)} güçlü yön, {len(AR_WEAKNESSES)} zayıf yön, {len(AR_OPPORTUNITIES)} fırsat, {len(AR_THREATS)} tehdit")
