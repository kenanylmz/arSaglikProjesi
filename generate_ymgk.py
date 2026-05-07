#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
YMGK Proje Takip Foyu - Nefes Saati
Kaynak: ymgk.docx
Cikti : YMGK_NefesiSaati.docx
"""

import shutil, copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\Kenan\Downloads\ymgk.docx'
DST = r'C:\Users\Kenan\Desktop\arSaglikProjesi\YMGK_NefesiSaati.docx'

# ── yardımcı fonksiyonlar ────────────────────────────────────────────────────

def _get_rpr(para):
    for r in para._p.findall(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None

def _make_rpr_plain():
    rpr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        fonts.set(qn(attr), 'Calibri')
    rpr.append(fonts)
    return rpr

def _make_run(text, rpr=None):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def _make_br_run(text, rpr=None):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    r.append(OxmlElement('w:br'))
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def _clear_runs(para):
    for tag in (qn('w:r'), qn('w:proofErr'),
                qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
        for elem in para._p.findall(tag):
            para._p.remove(elem)

def set_text(para, text):
    """Paragrafı temizleyip tek satır yazar."""
    rpr = _get_rpr(para) or _make_rpr_plain()
    for tag in (qn('w:b'), qn('w:bCs')):
        e = rpr.find(tag)
        if e is not None:
            rpr.remove(e)
    _clear_runs(para)
    if text:
        para._p.append(_make_run(text, rpr))

def set_lines(para, *lines):
    """Paragrafı temizleyip satırları w:br ile yazar."""
    rpr = _get_rpr(para) or _make_rpr_plain()
    for tag in (qn('w:b'), qn('w:bCs')):
        e = rpr.find(tag)
        if e is not None:
            rpr.remove(e)
    _clear_runs(para)
    for i, line in enumerate(lines):
        if i == 0:
            para._p.append(_make_run(line, rpr))
        else:
            para._p.append(_make_br_run(line, rpr))

def set_label(para, label, value):
    """'Label: değer' — label bold, değer normal."""
    rpr_plain = _get_rpr(para) or _make_rpr_plain()
    for tag in (qn('w:b'), qn('w:bCs')):
        e = rpr_plain.find(tag)
        if e is not None:
            rpr_plain.remove(e)
    rpr_bold = copy.deepcopy(rpr_plain)
    for tag in ('w:b', 'w:bCs'):
        rpr_bold.append(OxmlElement(tag))
    _clear_runs(para)
    para._p.append(_make_run(label, rpr_bold))
    para._p.append(_make_run(value, rpr_plain))

def set_cell(cell, *lines):
    """Hücreyi temizleyip satırlar halinde yazar."""
    para = cell.paragraphs[0]
    rpr  = _get_rpr(para) or _make_rpr_plain()
    for tag in (qn('w:b'), qn('w:bCs')):
        e = rpr.find(tag)
        if e is not None:
            rpr.remove(e)
    _clear_runs(para)
    for i, line in enumerate(lines):
        if i == 0:
            para._p.append(_make_run(line, rpr))
        else:
            para._p.append(_make_br_run(line, rpr))

# ── şablonu kopyala ──────────────────────────────────────────────────────────

shutil.copy(SRC, DST)
doc  = Document(DST)
p    = doc.paragraphs
t    = doc.tables

# ════════════════════════════════════════════════════════════════════════════
# 1. PROJE BİLGİLERİ
# ════════════════════════════════════════════════════════════════════════════

# [4] Proje Adı
set_label(p[4], 'Proje Adı : ',
    'Nefes Saati – AR Destekli Akıllı İlaç Takip ve Hatırlatma Sistemi')

# [5] Proje Konusu
set_lines(p[5],
    'Proje Konusu :',
    'Kronik hastalığa sahip bireyler ve yaşlı hastaların günlük ilaç uyumunu artırmak',
    'amacıyla geliştirilmiş Android tabanlı bir mobil sağlık uygulamasıdır. Artırılmış',
    'gerçeklik (AR) kamera ile ilaç kutusunu OCR teknolojisiyle otomatik tanır; saate özel',
    'günlük alarm bildirimleri gönderir; haftalık ilaç uyum takvimini görsel olarak sunar;',
    'bakıcılara AR-XXXX-XXXX koduyla kimlik doğrulamasız uzaktan izleme imkânı sağlar.')

# [6] Takım Adı + öğrenci bilgisi
set_lines(p[6],
    'Takım Adı : Nefes Saati',
    'Ad Soyad  : Kenan Yılmaz',
    'Okul No   : 210542011')

# ════════════════════════════════════════════════════════════════════════════
# 2. TAKIM ÜYELERİ — BOŞ BIRAK
# ════════════════════════════════════════════════════════════════════════════

# Örnek satırlarını temizle [9–12]
for i in [9, 10, 11, 12]:
    set_text(p[i], '')

# Tablo 0: başlık satırını koru, veri hücrelerini boşalt
for ri in range(1, len(t[0].rows)):
    for ci in range(len(t[0].columns)):
        set_cell(t[0].rows[ri].cells[ci], '')

# ════════════════════════════════════════════════════════════════════════════
# 3. PROJEDE KULLANILAN TEKNOLOJİLER  (Tablo 1)
# ════════════════════════════════════════════════════════════════════════════

tech = [
    ('React Native CLI 0.83 + TypeScript',
     'Android platformu cross-platform çatısı; TypeScript ile tip güvenli kod tabanı'),
    ('VisionCamera + Google ML Kit OCR',
     'AR kamera modülü; ilaç kutusu üzerinde gerçek zamanlı metin tanıma (OCR)'),
    ('Firebase Firestore + AsyncStorage + @notifee',
     'Local-first bulut senkronizasyonu; günlük tekrarlayan alarm bildirimleri'),
    ('Three.js + WebView + react-native-tts',
     '3D ilaç model görselleştirme; Türkçe sesli geri bildirim (TTS)'),
]

# Tablo 1'in veri satırları yetersizse satır ekle
while len(t[1].rows) - 1 < len(tech):
    t[1].add_row()

for i, (teknoloji, amac) in enumerate(tech):
    set_cell(t[1].rows[i + 1].cells[0], teknoloji)
    set_cell(t[1].rows[i + 1].cells[1], amac)

# ════════════════════════════════════════════════════════════════════════════
# 4. KARMA GERÇEKLİK (MR) KULLANIMI
# ════════════════════════════════════════════════════════════════════════════

# [20] Checkbox satırları — Telefon AR işaretli
set_lines(p[20],
    '☐ HoloLens',
    '☐ Meta Quest',
    '☑ Telefon AR',
    '☐ Unity AR Foundation',
    '☐ Diğer: ')

# [22] Açıklama
set_lines(p[22],
    'react-native-vision-camera ile Android kamerasının akışı Google ML Kit Vision API\'sine',
    'iletilir; OCR modülü ilaç kutusundaki metni gerçek zamanlı tanır ve kullanıcı onayına',
    'sunar. Three.js + WebView entegrasyonu ile GLTF formatındaki 3D ilaç modeli WebGL',
    'üzerinde render edilerek ilacın görsel temsili gösterilir.')

# [24] Kullanım Senaryosu
set_lines(p[24],
    'Ahmet Bey (68) uygulamayı açar, AR kamerasını ilaç kutusuna tutar; OCR ilaç adını',
    'otomatik tanır, onay ekranı açılır ve ilaç sisteme eklenir. Saat 08:00\'de alarm',
    'bildirimi gelir, "İlacı Aldım" butonuna basar; kayıt anlık Firebase\'e yazılır.',
    'Kızı Ayşe şehir dışından AR-XXXX-XXXX koduyla o günün doz durumunu takip eder.')

# ════════════════════════════════════════════════════════════════════════════
# 5. GITHUB REPO BİLGİLERİ
# ════════════════════════════════════════════════════════════════════════════

# [27] GitHub linki
set_lines(p[27],
    'GitHub Repo Linki :',
    'https://github.com/kenanylmz/arSaglikProjesi')

# [29] Repo içeriği checkboxları — hepsi işaretli
set_lines(p[29],
    '☑ Kaynak Kodlar',
    '☑ README Dosyası',
    '☑ Sprint Dokümanları')

# ════════════════════════════════════════════════════════════════════════════
# 6. TRELLO / GÖREV TAKİP PANOSU — KULLANILMADI
# ════════════════════════════════════════════════════════════════════════════

# [34] Trello linki
set_lines(p[34],
    'Trello / Jira / Notion Linki :',
    '(Bu projede Trello panosu kullanılmamıştır.)')

# Tablo 2 (Sprint tablosu) — Yapılacaklar / Devam / Tamamlananlar
sprint = [
    ('Firestore güvenlik\nkuralları sıkılaştırma',
     '—',
     'Faz 1: AR Kamera + OCR\nilaç tanıma modülü'),
    ('iOS platforma taşıma',
     '—',
     'Faz 2-3: @notifee alarmlar +\nFirebase sync + bakıcı modu'),
    ('Jest + Detox\ntest altyapısı',
     '—',
     'Faz 4-5: İmzalı APK +\nSWOT, RAMS, README, FOY2'),
]
for i, (yap, dev, tam) in enumerate(sprint):
    set_cell(t[2].rows[i + 1].cells[0], yap)
    set_cell(t[2].rows[i + 1].cells[1], dev)
    set_cell(t[2].rows[i + 1].cells[2], tam)

# ════════════════════════════════════════════════════════════════════════════
# 7. SWOT ANALİZİ
# ════════════════════════════════════════════════════════════════════════════

# Tablo 3 — Güçlü / Zayıf
set_cell(t[3].rows[1].cells[0],
    '• AR+OCR ile otomatik ilaç kutusu tanıma',
    '• Local-first + Auth-free mimari (%100 çevrimdışı)',
    '• Bakıcı uzaktan izleme (AR-XXXX-XXXX kodu)')
set_cell(t[3].rows[1].cells[1],
    '• Yalnızca Android platformu (iOS yok)',
    '• Açık Firestore güvenlik kuralları',
    '• OCR tanıma güvenilirlik riski')
set_cell(t[3].rows[2].cells[0],
    '• Türkçe TTS erişilebilirliği',
    '• "İlacı Aldım" onay + haftalık uyum takvimi',
    '• Günlük akıllı alarm (23:00–07:00 sessiz)')
set_cell(t[3].rows[2].cells[1],
    '• Sınırlı ilaç veritabanı',
    '• Bakıcı modu tek yönlü (read-only)',
    '• İlaç etkileşim bilgisi yok')

# Tablo 4 — Fırsatlar / Tehditler
set_cell(t[4].rows[1].cells[0],
    '• Hastane taburculuk paketine entegrasyon',
    '• Firebase → doktor portal genişlemesi',
    '• iOS taşıma ve IoT ilaç kutusu entegrasyonu')
set_cell(t[4].rows[1].cells[1],
    '• Açık Firestore kurallarından güvenlik riski',
    '• TİTCK / KVKK uyumluluk yükümlülükleri',
    '• Medisafe/MyTherapy gibi olgun rakipler')
set_cell(t[4].rows[2].cells[0],
    '• ML tabanlı görüntü tanıma ile OCR geliştirme',
    '• Akademik araştırma için Firestore uyum verisi',
    '• Giyilebilir cihaz (Wear OS) bildirim desteği')
set_cell(t[4].rows[2].cells[1],
    '• Firebase vendor bağımlılığı riski',
    '• React Native büyük sürüm kırılma riski',
    '• Yaşlı kullanıcı adaptasyon zorluğu')

# ════════════════════════════════════════════════════════════════════════════
# 8. PROJE GELİŞİM DURUMU  (Tablo 5)
# ════════════════════════════════════════════════════════════════════════════

set_cell(t[5].rows[1].cells[0],
    'Faz 1 TAMAMLANDI: VisionCamera + ML Kit OCR ilaç kutusu tanıma modülü,',
    '3D ilaç model görselleştirme (Three.js + WebView + GLTFLoader), Türkçe TTS.')
set_cell(t[5].rows[2].cells[0],
    'Faz 2-3 TAMAMLANDI: @notifee günlük alarm sistemi, Firebase Firestore',
    'senkronizasyonu, AR-XXXX-XXXX anonim kimlik, bakıcı uzaktan izleme modu,',
    '"İlacı Aldım" onay akışı, 7 günlük haftalık uyum takvimi.')
set_cell(t[5].rows[3].cells[0],
    'Faz 4-5 TAMAMLANDI: Gradle + PKCS12 Keystore imzalı APK üretimi;',
    'PROJE_DOKUMANI.md (5 faz), SWOT Analizi, RAMS Analizi, README.md,',
    'FÖY2 ve Genel Döküman belgelerinin tamamı hazırlanmıştır.')

# ════════════════════════════════════════════════════════════════════════════
# 9. KARŞILAŞILAN PROBLEMLER
# ════════════════════════════════════════════════════════════════════════════

set_lines(p[45],
    '1. TypeScript navigasyon tip uyumsuzluğu → Bottom tab mimarisine geçişle çözüldü.',
    '2. Gradle kilit hatası (Java süreci çakışması) → Stop-Process ile sonlandırıldı.',
    '3. ic_launcher_round eksikliği → AndroidManifest roundIcon ic_launcher olarak güncellendi.',
    '4. keytool PATH sorunu (JDK-19) → Tam yol ile çağrıldı.',
    '5. Auth-free Firestore veri izolasyonu → AR-XXXX-XXXX anonim ID mimarisi geliştirildi.')

# ════════════════════════════════════════════════════════════════════════════
# 10. PLANLANAN SONRAKİ ADIMLAR
# ════════════════════════════════════════════════════════════════════════════

set_lines(p[48],
    '1. Firestore güvenlik kurallarının üretim ortamı için sıkılaştırılması (kritik).',
    '2. iOS platforma taşıma (react-native-vision-camera cross-platform desteği mevcut).',
    '3. Otomatik test altyapısı: Jest birim testleri + Detox E2E testleri.',
    '4. Bakıcı modu çift yönlü iletişim ve React/Next.js doktor portal entegrasyonu.')

# ── kaydet ───────────────────────────────────────────────────────────────────

doc.save(DST)
print(f'Belge başarıyla oluşturuldu: {DST}')
