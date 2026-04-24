#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
FOY2 - Nefes Saati projesi icin doldurulmus Word belgesi olusturur.
Kaynak sablon: FÖY2.docx
Cikti: FOY2_NefesiSaati.docx
"""

import shutil
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\Kenan\Downloads\FÖY2.docx'
DST = r'C:\Users\Kenan\Desktop\arSaglikProjesi\FOY2_NefesiSaati.docx'

# ── yardımcı fonksiyonlar ────────────────────────────────────────────────────

def _get_rpr(para):
    """Paragraftaki ilk run'ın rPr elementinin derin kopyasını döner."""
    for r in para._p.findall(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None

def _make_run(text, rpr=None):
    """Metin ve opsiyonel rPr ile w:r elementi oluşturur."""
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def _clear_runs(para):
    """Paragraftaki tüm run, proofErr, bookmark elementlerini temizler."""
    p = para._p
    for tag in (qn('w:r'), qn('w:proofErr'),
                qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
        for elem in p.findall(tag):
            p.remove(elem)

def set_text(para, text, bold=None):
    """Paragrafın içeriğini değiştirir, pPr (paragraf biçimi) korunur."""
    rpr = _get_rpr(para)
    _clear_runs(para)
    if rpr is not None and bold is not None:
        for tag in (qn('w:b'), qn('w:bCs')):
            existing = rpr.find(tag)
            if bold and existing is None:
                rpr.append(OxmlElement(tag))
            elif not bold and existing is not None:
                rpr.remove(existing)
    para._p.append(_make_run(text, rpr))

def append_text(para, text):
    """Mevcut paragrafın sonuna kalın-olmayan yeni bir run ekler."""
    rpr = _get_rpr(para)
    if rpr is not None:
        for tag in (qn('w:b'), qn('w:bCs')):
            existing = rpr.find(tag)
            if existing is not None:
                rpr.remove(existing)
    para._p.append(_make_run(text, rpr))

def set_cell(cell, text):
    """Tablo hücresinin metnini değiştirir."""
    para = cell.paragraphs[0]
    rpr = _get_rpr(para)
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            fonts.set(qn(attr), 'Times New Roman')
        kern = OxmlElement('w:kern')
        kern.set(qn('w:val'), '0')
        rpr.append(fonts)
        rpr.append(kern)
    _clear_runs(para)
    para._p.append(_make_run(text, rpr))

# ── şablonu kopyala ve aç ───────────────────────────────────────────────────

shutil.copy(SRC, DST)
doc = Document(DST)
paras = doc.paragraphs
table = doc.tables[0]

# ── öğrenci bilgileri (para 1-2) ─────────────────────────────────────────────

set_text(paras[1], 'Ad Soyad: Kenan Yılmaz')
set_text(paras[2], 'Numara: 210542011')

# ── karma gerçeklik açıklaması (para 4) ──────────────────────────────────────

karma_aciklama = (
    "Karma Gerçeklik (Mixed Reality – MR), gerçek dünya ile dijital/sanal "
    "nesnelerin bir arada bulunduğu ve birbirleriyle etkileşime girebildiği "
    "teknolojidir. Artırılmış Gerçeklik (AR) ile Sanal Gerçeklik (VR) "
    "spektrumunun ortasında yer alan MR'de sanal nesneler gerçek dünyayı "
    "tanıyarak fiziksel nesnelerle etkileşime girebilir. Bu sayede holografik "
    "görüntüleme, 3D model yerleştirme ve gerçek zamanlı fiziksel–dijital "
    "etkileşim mümkün hâle gelir. Sağlık, eğitim, endüstri ve mimarlık gibi "
    "alanlarda büyük potansiyel taşıyan MR teknolojisi; ilaç tanıma, cerrahi "
    "planlama ve tıbbi eğitim gibi kritik sağlık uygulamalarında giderek artan "
    "bir kullanım alanı bulmaktadır."
)
set_text(paras[4], karma_aciklama)
set_text(paras[5], '')  # boş devam satırını temizle

# ── MR ekosistem tablosu ─────────────────────────────────────────────────────

set_cell(table.rows[1].cells[1],
    'Gerçek dünyayı algılamak ve görüntülemek için kullanılan fiziksel cihazlar; '
    'kameralar, sensörler, işlemci ve ekran birimlerini kapsar.')
set_cell(table.rows[1].cells[2],
    'ARCore destekli Android akıllı telefon, derinlik sensörü, '
    'çift kamera sistemi, giyilebilir AR gözlük (HoloLens)')

set_cell(table.rows[2].cells[1],
    'AR/MR deneyimini oluşturan platform, geliştirme çatısı ve kütüphane '
    'katmanları; görüntü işleme, sahne yönetimi ve kullanıcı arayüzünü kapsar.')
set_cell(table.rows[2].cells[2],
    'React Native CLI, VisionCamera, Google ML Kit (OCR), '
    'Three.js/WebGL (3D render), Firebase Firestore, @notifee')

set_cell(table.rows[3].cells[1],
    'Kullanıcının sanal ve gerçek dünya nesneleriyle etkileşimini sağlayan '
    'giriş/çıkış yöntemleri ve geri bildirim mekanizmaları.')
set_cell(table.rows[3].cells[2],
    'Dokunmatik ekran, kamera hareketi tanıma, OCR ile ilaç kutusu tanıma, '
    'Türkçe sesli geri bildirim (TTS), günlük push bildirim alarmları')

# ── proje başlığı (para 10) ──────────────────────────────────────────────────

set_text(paras[10],
    'Nefes Saati – AR Destekli Akıllı İlaç Takip ve Hatırlatma Sistemi')

# ── problem tanımı (para 12) ─────────────────────────────────────────────────

problem = (
    "Kronik hastalığa sahip bireyler ve yaşlı hastalar, günlük ilaç dozlarını "
    "düzenli olarak almakta güçlük çekmektedir. Hangi ilacın ne zaman alınacağını "
    "unutmak, yanlış doz ve çoklu ilaç karışıklığı ciddi sağlık risklerine yol "
    "açmaktadır. Aile üyeleri ve bakıcılar ise hastanın ilaç uyumunu gerçek zamanlı "
    "olarak uzaktan takip edememektedir."
)
set_text(paras[12], problem)

# ── önerilen çözüm (para 14) ─────────────────────────────────────────────────

cozum = (
    "React Native CLI 0.83 ile geliştirilen Nefes Saati uygulaması; VisionCamera "
    "ve Google ML Kit OCR ile ilaç kutusundaki yazıları tanıyarak ilacı otomatik "
    "olarak sisteme ekler. @notifee kütüphanesiyle günlük, saate özel alarm "
    "bildirimleri gönderir. Firebase Firestore üzerinden haftalık doz geçmişi buluta "
    "senkronize edilir. Kullanıcı 'İlacı Aldım' butonuyla doz onayı verir; bakıcılar "
    "ise AR-XXXX-XXXX kimlik koduyla yakınlarının ilaç uyumunu anlık izleyebilir."
)
set_text(paras[14], cozum)

# ── proje beklentileri – soru yanıtları ──────────────────────────────────────

append_text(paras[17],
    'Kronik hastalığa sahip bireylerin ve yaşlı hastaların ilaç dozlarını '
    'düzenli almasını desteklemek; aile üyeleri ve bakıcıların uzaktan ilaç '
    'uyumunu izleyebilmesini ve kişisel sağlık yönetiminin dijitalleştirilmesini '
    'sağlamaktadır.')

append_text(paras[18],
    'Kullanıcılar AR kamera ile ilaçlarını hızlıca sisteme ekleyebilir, '
    'otomatik bildirimler sayesinde zamanında hatırlatma alabilir ve '
    'haftalık ilaç uyum istatistiklerini takip edebilir. Bakıcılar ise '
    'anlık doz durumuna ve alım geçmişine uzaktan erişebilir.')

append_text(paras[19],
    ' OCR tabanlı AR ilaç tanıma, kimlik doğrulama gerektirmeyen benzersiz '
    'cihaz kimliği (AR-XXXX-XXXX) ile bakıcı izleme modu ve çevrimdışı '
    'çalışabilen yerel-önce (local-first) mimarisinin bütünleşik bir sağlık '
    'takip ekosistemi hâlinde sunulmasıdır.')

# ── kaydet ───────────────────────────────────────────────────────────────────

doc.save(DST)
print(f'Belge başarıyla oluşturuldu: {DST}')
